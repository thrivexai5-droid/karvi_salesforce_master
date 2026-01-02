from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.contrib.auth.models import User
from django import forms
from django.core.paginator import Paginator
from django.db.models import Q, Sum
from django.db import models
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from datetime import datetime
import json
import base64
import os
from .models import UserProfile, Company, Contact, PurchaseOrder, PurchaseOrderItem, Invoice, InquiryHandler, InquiryItem, Quotation, AdditionalSupply, Notification
from .password_storage import password_storage

def increment_revision(current_revision):
    """
    Auto-increment revision from Rev A to Rev B, Rev C, etc.
    """
    if not current_revision:
        return 'Rev A'
    
    if not current_revision.startswith('Rev '):
        return 'Rev B'  # If it doesn't start with 'Rev ', assume it's been modified and increment to B
    
    try:
        # Extract the letter part (A, B, C, etc.)
        revision_letter = current_revision.split('Rev ')[1].strip()
        
        if len(revision_letter) == 1 and revision_letter.isalpha():
            # Single letter revision (A, B, C, etc.)
            next_letter = chr(ord(revision_letter.upper()) + 1)
            return f'Rev {next_letter}'
        else:
            # If it's not a single letter, default to Rev B
            return 'Rev B'
    except (IndexError, ValueError):
        # If parsing fails, default to Rev B
        return 'Rev B'

# Import Mistral AI for PDF processing
try:
    from mistralai import Mistral
    MISTRAL_AVAILABLE = True
except ImportError:
    MISTRAL_AVAILABLE = False

@login_required
def debug_dashboard_data(request):
    """Debug endpoint to check dashboard data"""
    try:
        # Check purchase orders
        po_count = PurchaseOrder.objects.count()
        recent_pos = PurchaseOrder.objects.select_related('company__company').order_by('-order_date', '-id')[:5]
        
        # Check invoices
        invoice_count = Invoice.objects.count()
        recent_invoices = Invoice.objects.select_related('company__company').order_by('-invoice_date', '-id')[:5]
        
        # Check user role
        user_role = request.user.userprofile.get_roles_list() if hasattr(request.user, 'userprofile') else 'No profile'
        
        debug_data = {
            'user': request.user.username,
            'user_role': user_role,
            'purchase_orders': {
                'count': po_count,
                'recent': [
                    {
                        'po_number': po.po_number,
                        'customer_name': po.customer_name,
                        'company': po.company.company.company_name if po.company.company else po.company.customer_name,
                        'order_date': po.order_date.strftime('%Y-%m-%d'),
                        'due_days': po.due_days,
                        'sales_person': po.sales_person.username if po.sales_person else None
                    } for po in recent_pos
                ]
            },
            'invoices': {
                'count': invoice_count,
                'recent': [
                    {
                        'invoice_number': inv.invoice_number,
                        'company': inv.company.company.company_name if inv.company.company else inv.company.customer_name,
                        'invoice_date': inv.invoice_date.strftime('%Y-%m-%d')
                    } for inv in recent_invoices
                ]
            }
        }
        
        return JsonResponse({
            'success': True,
            'data': debug_data
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })

@login_required
def test_mistral_connection(request):
    """Test endpoint to verify Mistral API connection"""
    try:
        from mistralai import Mistral
        from django.conf import settings
        
        # Check if API key is configured
        if not settings.MISTRAL_API_KEY or settings.MISTRAL_API_KEY == 'your-actual-mistral-api-key-here':
            return JsonResponse({
                'success': False,
                'error': 'Mistral API key is not configured properly'
            })
        
        # Test simple API call
        client = Mistral(api_key=settings.MISTRAL_API_KEY)
        
        # Simple test call
        response = client.chat.complete(
            model="mistral-medium-latest",
            messages=[
                {"role": "user", "content": "Hello, this is a test. Please respond with 'API connection successful'."}
            ],
            temperature=0
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Mistral API connection successful',
            'response': response.choices[0].message.content
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Mistral API connection failed: {str(e)}'
        })

@login_required
@csrf_exempt
def upload_po_ajax(request):
    """AJAX endpoint that receives the PDF, calls the AI service, and returns JSON data to the frontend."""
    if request.method == 'POST' and request.FILES.get('po_file'):
        try:
            # Import the service function
            from .services import extract_po_data_from_pdf
            
            # Get the uploaded file
            pdf_file = request.FILES['po_file']
            
            # Validate file type
            if not pdf_file.name.lower().endswith('.pdf'):
                return JsonResponse({
                    'success': False,
                    'error': 'Please upload a PDF file only.'
                })
            
            # Validate file size (5MB limit)
            if pdf_file.size > 5 * 1024 * 1024:
                return JsonResponse({
                    'success': False,
                    'error': 'File size must be less than 5MB.'
                })
            
            # Call the service
            extracted_data = extract_po_data_from_pdf(pdf_file)
            
            # Check if extraction failed
            if extracted_data.get('error'):
                return JsonResponse({
                    'success': False,
                    'error': extracted_data['error']
                })
            
            # Log extracted data for debugging (optional)
            import logging
            logger = logging.getLogger(__name__)
            logger.info(f"Extracted PO data: {extracted_data}")
            
            return JsonResponse({'success': True, 'data': extracted_data})
            
        except ImportError as e:
            return JsonResponse({
                'success': False, 
                'error': f'Service import error: {str(e)}'
            })
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            print(f"AI Extraction Error: {error_details}")  # For debugging
            return JsonResponse({
                'success': False, 
                'error': f'Error processing PDF: {str(e)}'
            })
    
    return JsonResponse({'success': False, 'error': 'No file provided.'})

@login_required
@csrf_exempt
def save_purchase_order_items_ajax(request):
    """AJAX endpoint to save purchase order items"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            purchase_order_id = data.get('purchase_order_id')
            items_data = data.get('items', [])
            
            if not purchase_order_id:
                return JsonResponse({'success': False, 'error': 'Purchase order ID is required'})
            
            # Get the purchase order
            try:
                purchase_order = PurchaseOrder.objects.get(id=purchase_order_id)
            except PurchaseOrder.DoesNotExist:
                return JsonResponse({'success': False, 'error': 'Purchase order not found'})
            
            # Clear existing items
            purchase_order.items.all().delete()
            
            # Add new items
            total_amount = 0
            for item_data in items_data:
                if item_data.get('item_name') and item_data.get('quantity') and item_data.get('price'):
                    # Extract material code from item name if it's in [CODE] format
                    item_name = item_data['item_name']
                    material_code = ''
                    
                    if item_name.startswith('[') and ']' in item_name:
                        parts = item_name.split(']', 1)
                        material_code = parts[0][1:]  # Remove the opening bracket
                        item_name = parts[1].strip() if len(parts) > 1 else item_name
                    
                    item = PurchaseOrderItem.objects.create(
                        purchase_order=purchase_order,
                        material_code=material_code,
                        item_name=item_name,
                        quantity=float(item_data['quantity']),
                        price=float(item_data['price'])
                    )
                    total_amount += float(item.amount)
            
            # Update purchase order total value
            purchase_order.order_value = total_amount
            purchase_order.save()
            
            return JsonResponse({
                'success': True,
                'total_amount': total_amount,
                'message': f'Saved {len(items_data)} items successfully'
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method'})

@login_required
def get_purchase_order_items_ajax(request):
    """AJAX endpoint to get purchase order items"""
    purchase_order_id = request.GET.get('purchase_order_id')
    
    print(f"DEBUG: get_purchase_order_items_ajax called with PO ID: {purchase_order_id}")
    
    if not purchase_order_id:
        return JsonResponse({'success': False, 'error': 'Purchase order ID is required'})
    
    try:
        purchase_order = PurchaseOrder.objects.get(id=purchase_order_id)
        items = purchase_order.items.all()
        
        print(f"DEBUG: Found {items.count()} items for PO {purchase_order.po_number}")
        
        items_data = []
        total_amount = 0
        
        for item in items:
            # Combine material code and item name for display
            display_name = item.item_name
            if hasattr(item, 'material_code') and item.material_code:
                display_name = f"[{item.material_code}] {item.item_name}"
            
            item_data = {
                'id': item.id,
                'item_name': display_name,
                'quantity': str(item.quantity),
                'price': str(item.price),
                'amount': str(item.amount)
            }
            items_data.append(item_data)
            total_amount += float(item.amount)
            print(f"DEBUG: Item data: {item_data}")
        
        response_data = {
            'success': True,
            'items': items_data,
            'total_amount': total_amount
        }
        print(f"DEBUG: Returning response: {response_data}")
        return JsonResponse(response_data)
        
    except PurchaseOrder.DoesNotExist:
        print(f"DEBUG: Purchase order {purchase_order_id} not found")
        return JsonResponse({'success': False, 'error': 'Purchase order not found'})
    except Exception as e:
        print(f"DEBUG: Error in get_purchase_order_items_ajax: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
def search_customers_ajax(request):
    """AJAX endpoint to search customers for autocomplete"""
    query = request.GET.get('q', '').strip()
    
    if len(query) < 2:  # Require at least 2 characters
        return JsonResponse({'success': True, 'customers': []})
    
    try:
        # Search in customer_name and contact_name fields
        customers = Contact.objects.filter(
            models.Q(customer_name__icontains=query) |
            models.Q(contact_name__icontains=query)
        ).select_related('company').order_by('customer_name')[:10]  # Limit to 10 results
        
        customers_data = []
        for customer in customers:
            customers_data.append({
                'id': customer.id,
                'customer_name': customer.customer_name,
                'contact_name': customer.contact_name,
                'company_name': customer.company.company_name if customer.company else '',
                'display_text': f"{customer.customer_name} - {customer.company.company_name if customer.company else 'No Company'}"
            })
        
        return JsonResponse({
            'success': True,
            'customers': customers_data
        })
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

class CustomUserCreationForm(UserCreationForm):
    email = forms.EmailField(required=True)
    
    class Meta:
        model = User
        fields = ("username", "email", "password1", "password2")
    
    def save(self, commit=True):
        user = super().save(commit=False)
        user.email = self.cleaned_data["email"]
        if commit:
            user.save()
        return user

class UserManagementForm(forms.Form):
    # Basic Information
    name = forms.CharField(
        max_length=100, 
        required=True,
        widget=forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Enter full name'})
    )
    email = forms.EmailField(
        required=True,
        widget=forms.EmailInput(attrs={'class': 'form-control', 'placeholder': 'Enter email address'})
    )
    phone_number = forms.CharField(
        max_length=20,
        required=False,
        widget=forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Enter phone number'})
    )
    password = forms.CharField(
        min_length=8,
        required=True,
        widget=forms.PasswordInput(attrs={'class': 'form-control', 'placeholder': 'Enter password'})
    )
    
    # Role Selection
    ROLE_CHOICES = [
        ('sales', 'Sales'),
        ('project_manager', 'Project Manager'),
    ]
    role = forms.ChoiceField(
        choices=ROLE_CHOICES,
        required=True,
        widget=forms.Select(attrs={'class': 'form-select'})
    )
    
    # Form Permissions - Multiple Select
    FORM_CHOICES = [
        ('invoice_generation', 'Invoice Generation'),
        ('inquiry_handler', 'Inquiry Handler'),
        ('quotation_generation', 'Quotation Generation'),
        ('additional_supply', 'Additional Supply'),
    ]
    
    form_permissions = forms.MultipleChoiceField(
        choices=FORM_CHOICES,
        required=False,
        widget=forms.CheckboxSelectMultiple(attrs={'class': 'form-check-input'}),
        help_text="Select multiple forms this user can access"
    )
    
    def clean_email(self):
        email = self.cleaned_data.get('email')
        if User.objects.filter(email=email).exists():
            raise forms.ValidationError("A user with this email already exists.")
        return email
    
    def clean(self):
        cleaned_data = super().clean()
        role = cleaned_data.get('role')
        form_permissions = cleaned_data.get('form_permissions', [])
        
        # Role-based permission validation
        if role == 'project_manager' and form_permissions:
            invalid_perms = [perm for perm in form_permissions if perm != 'additional_supply']
            if invalid_perms:
                raise forms.ValidationError("Project Managers can only access Additional Supply.")
        
        if role == 'sales' and form_permissions:
            allowed_perms = ['invoice_generation', 'inquiry_handler', 'quotation_generation']
            invalid_perms = [perm for perm in form_permissions if perm not in allowed_perms]
            if invalid_perms:
                raise forms.ValidationError("Sales users can only access Invoice Generation, Inquiry Handler, or Quotation Generation.")
        
        return cleaned_data

class UserEditForm(forms.Form):
    # Basic Information
    name = forms.CharField(
        max_length=100, 
        required=True,
        widget=forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Enter full name'})
    )
    email = forms.EmailField(
        required=True,
        widget=forms.EmailInput(attrs={'class': 'form-control', 'placeholder': 'Enter email address'})
    )
    phone_number = forms.CharField(
        max_length=20,
        required=False,
        widget=forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Enter phone number'})
    )
    password = forms.CharField(
        min_length=8,
        required=False,  # Optional for edit
        widget=forms.PasswordInput(attrs={'class': 'form-control', 'placeholder': 'Leave blank to keep current password'})
    )
    
    # Role Selection
    ROLE_CHOICES = [
        ('sales', 'Sales'),
        ('project_manager', 'Project Manager'),
    ]
    role = forms.ChoiceField(
        choices=ROLE_CHOICES,
        required=True,
        widget=forms.Select(attrs={'class': 'form-select'})
    )
    
    # Form Permissions - Multiple Select
    FORM_CHOICES = [
        ('invoice_generation', 'Invoice Generation'),
        ('inquiry_handler', 'Inquiry Handler'),
        ('quotation_generation', 'Quotation Generation'),
        ('additional_supply', 'Additional Supply'),
    ]
    
    form_permissions = forms.MultipleChoiceField(
        choices=FORM_CHOICES,
        required=False,
        widget=forms.CheckboxSelectMultiple(attrs={'class': 'form-check-input'}),
        help_text="Select multiple forms this user can access"
    )
    
    def __init__(self, *args, **kwargs):
        self.user_instance = kwargs.pop('user_instance', None)
        super().__init__(*args, **kwargs)
    
    def clean_email(self):
        email = self.cleaned_data.get('email')
        if self.user_instance and User.objects.filter(email=email).exclude(id=self.user_instance.id).exists():
            raise forms.ValidationError("A user with this email already exists.")
        elif not self.user_instance and User.objects.filter(email=email).exists():
            raise forms.ValidationError("A user with this email already exists.")
        return email
    
    def clean(self):
        cleaned_data = super().clean()
        role = cleaned_data.get('role')
        form_permissions = cleaned_data.get('form_permissions', [])
        
        # Role-based permission validation
        if role == 'project_manager' and form_permissions:
            invalid_perms = [perm for perm in form_permissions if perm != 'additional_supply']
            if invalid_perms:
                raise forms.ValidationError("Project Managers can only access Additional Supply.")
        
        if role == 'sales' and form_permissions:
            allowed_perms = ['invoice_generation', 'inquiry_handler', 'quotation_generation']
            invalid_perms = [perm for perm in form_permissions if perm not in allowed_perms]
            if invalid_perms:
                raise forms.ValidationError("Sales users can only access Invoice Generation, Inquiry Handler, or Quotation Generation.")
        
        return cleaned_data

@login_required
def dashboard_view(request):
    """Main dashboard view with role-based content - requires authentication"""
    from decimal import Decimal
    from django.db.models import Count, Case, When, IntegerField, Max
    from datetime import datetime
    
    # Get user role
    user_role = request.user.userprofile.get_roles_list()
    
    # Base context for all users
    context = {
        'user_role': user_role,
    }
    
    # Role-specific dashboard content
    if user_role in ['admin', 'manager']:
        # Admin/Manager - Full dashboard with all metrics
        context.update(get_full_dashboard_data())
    elif user_role == 'sales':
        # Sales - Focus on inquiries, quotations, and invoices (user-specific)
        context.update(get_sales_dashboard_data(request.user))
    elif user_role == 'project_manager':
        # Project Manager - Focus on additional supplies and project tracking
        context.update(get_project_manager_dashboard_data(request.user))
    else:
        # Default - Basic dashboard
        context.update(get_basic_dashboard_data())
    
    return render(request, 'dashboard/index.html', context)


def get_full_dashboard_data():
    """Get complete dashboard data for admin/manager roles"""
    from decimal import Decimal
    from django.db.models import Count, Case, When, IntegerField, Max
    from datetime import datetime
    
    # Calculate dynamic values from Invoice model
    total_value_result = Invoice.objects.aggregate(total=Sum('order_value'))
    total_value = total_value_result['total'] or Decimal('0')
    
    # Calculate GST as 18% of Total Value (using Decimal for precision)
    gst_rate = Decimal('0.18')
    gst_value = total_value * gst_rate
    
    # Calculate Max Date - Latest payment due date from invoices
    max_payment_due_date = Invoice.objects.aggregate(max_date=Max('payment_due_date'))['max_date']
    max_invoice_date = Invoice.objects.aggregate(max_date=Max('invoice_date'))['max_date']
    max_inquiry_date = InquiryHandler.objects.aggregate(max_date=Max('date_of_quote'))['max_date']
    max_po_date = PurchaseOrder.objects.aggregate(max_date=Max('order_date'))['max_date']
    
    # Use payment due date as priority, fallback to other dates
    dates_list = [d for d in [max_payment_due_date, max_invoice_date, max_inquiry_date, max_po_date] if d is not None]
    max_date = max(dates_list) if dates_list else None
    
    # Format max date for display (DD-MM-YYYY)
    max_date_formatted = max_date.strftime('%d-%m-%Y') if max_date else '—'
    
    # Calculate Sustainance Date - Using the latest created_at date as system-defined date
    from django.utils import timezone
    
    # Get timezone-aware datetime.min for comparison
    timezone_aware_min = timezone.make_aware(datetime.min) if timezone.is_naive(datetime.min) else datetime.min
    
    max_created_date = max(
        Invoice.objects.aggregate(max_created=Max('created_at'))['max_created'] or timezone_aware_min,
        InquiryHandler.objects.aggregate(max_created=Max('created_at'))['max_created'] or timezone_aware_min,
        PurchaseOrder.objects.aggregate(max_created=Max('created_at'))['max_created'] or timezone_aware_min
    )
    
    # Format sustainance date for display (DD-MM-YYYY)
    sustainance_date_formatted = max_created_date.strftime('%d-%m-%Y') if max_created_date != timezone_aware_min else '—'
    
    # Format values for display (Indian number format)
    def format_indian_currency(amount):
        """Format number in Indian currency style (e.g., 4,21,50,000)"""
        if amount == 0:
            return "0"
        
        # Convert to float for formatting, then back to string
        amount_float = float(amount)
        amount_str = f"{amount_float:,.0f}"  # No decimal places for currency display
        
        return amount_str
    
    # Get dynamic inquiry status data for Leads chart
    def get_inquiry_status_data():
        """Fetch inquiry counts by status for Leads chart"""
        import json
        
        # Define the exact statuses to include in the chart (as per requirements)
        chart_statuses = [
            'Inputs', 'Pending', 'Inspection', 'Inquiry', 'Quotation', 'Negotiation',
            'Inquiry Hold', 'PO-Confirm', 'Design Review', 'Manufacturing',
            'Stage-Inspection', 'Dispatch', 'GRN', 'Project Closed', 'Lost', 'PO Hold'
        ]
        
        # Get counts for each status using a single optimized query
        status_counts = InquiryHandler.objects.aggregate(
            **{
                status.lower().replace('-', '_').replace(' ', '_'): Count(
                    Case(When(status=status, then=1), output_field=IntegerField())
                )
                for status in chart_statuses
            }
        )
        
        # Format data for chart consumption
        data_values = [status_counts.get(status.lower().replace('-', '_').replace(' ', '_'), 0) for status in chart_statuses]
        
        chart_data = {
            'categories': json.dumps(chart_statuses),
            'data': json.dumps(data_values)
        }
        
        return chart_data
    
    # Get inquiry status data
    inquiry_chart_data = get_inquiry_status_data()
    
    # Calculate paid invoice data for Operations and Collection values
    paid_total_value = Invoice.objects.filter(status='paid').aggregate(total=Sum('order_value'))['total'] or Decimal('0')
    paid_invoices_count = Invoice.objects.filter(status='paid').count()
    
    # Calculate sustainability metrics
    from .services import calculate_sustainability_date
    sustainability_data = calculate_sustainability_date()
    
    # Handle negative sustainability case
    if sustainability_data['sustainability_days'] <= 0:
        sustainability_status = "Critical"
        sustainability_color = "danger"
    elif sustainability_data['sustainability_days'] <= 30:
        sustainability_status = "Warning"
        sustainability_color = "warning"
    else:
        sustainability_status = "Healthy"
        sustainability_color = "success"
    
    # Get recent invoices for the dashboard card (latest 10)
    recent_invoices = Invoice.objects.select_related('company__company').order_by('-invoice_date', '-id')[:10]
    
    # Get recent purchase orders (all POs for admin/manager)
    recent_purchase_orders = PurchaseOrder.objects.select_related('company__company').order_by('-order_date', '-id')[:10]
    total_purchase_orders = PurchaseOrder.objects.count()
    
    # Get recent quotations (all quotations for admin/manager) - using InquiryHandler for running projects
    recent_quotations = InquiryHandler.objects.select_related('company__company').exclude(
        status__in=['Project Closed', 'Lost']
    ).order_by('-date_of_quote')[:10]
    total_quotations = InquiryHandler.objects.exclude(
        status__in=['Project Closed', 'Lost']
    ).count()
    
    return {
        'dashboard_type': 'full',
        'max_date': max_date_formatted,
        'max_date_label': 'Payment Due (MAX)',
        'sustainance_date': sustainability_data['sustainability_date'].strftime('%d-%m-%Y'),
        'sustainability_days': sustainability_data['sustainability_days'],
        'sustainability_status': sustainability_status,
        'sustainability_color': sustainability_color,
        'net_revenue': format_indian_currency(abs(sustainability_data['net_revenue'])),  # Show absolute value
        'net_revenue_sign': 'positive' if sustainability_data['net_revenue'] >= 0 else 'negative',
        'monthly_expenses': format_indian_currency(sustainability_data['monthly_expenses']),
        'daily_burn': format_indian_currency(sustainability_data['daily_burn']),
        'total_value': format_indian_currency(total_value),
        'gst_value': format_indian_currency(gst_value),
        'total_value_raw': total_value,
        'gst_value_raw': gst_value,
        'paid_total_value': format_indian_currency(paid_total_value),
        'paid_total_value_raw': paid_total_value,
        'paid_invoices_count': paid_invoices_count,
        'recent_invoices': recent_invoices,
        'recent_purchase_orders': recent_purchase_orders,
        'total_purchase_orders': total_purchase_orders,
        'recent_quotations': recent_quotations,
        'total_quotations': total_quotations,
        'is_user_specific': False,
        'inquiry_chart_data': inquiry_chart_data,
    }


def get_sales_dashboard_data(user=None):
    """Get sales-focused dashboard data - filtered by user if provided"""
    from decimal import Decimal
    from django.db.models import Sum, F, DecimalField
    from django.db.models.functions import Coalesce
    import json
    
    def format_indian_currency(amount):
        if amount == 0:
            return "0"
        amount_float = float(amount)
        return f"{amount_float:,.0f}"
    
    # Filter inquiries by user if provided (for sales users)
    if user and user.userprofile.get_roles_list() == 'sales':
        # Sales user - only their assigned inquiries and POs
        inquiry_filter = {'sales': user}
        po_filter = {'sales_person': user}
        total_inquiries = InquiryHandler.objects.filter(sales=user).count()
        active_inquiries = InquiryHandler.objects.filter(sales=user).exclude(status__in=['Project Closed', 'Lost']).count()
        quotations_sent = InquiryHandler.objects.filter(sales=user, status='Quotation').count()
        is_user_specific = True
    else:
        # Admin or no user specified - all inquiries and POs
        inquiry_filter = {}
        po_filter = {}
        total_inquiries = InquiryHandler.objects.count()
        active_inquiries = InquiryHandler.objects.exclude(status__in=['Project Closed', 'Lost']).count()
        quotations_sent = InquiryHandler.objects.filter(status='Quotation').count()
        is_user_specific = False
    
    # Invoice metrics for sales (all invoices for now - can be filtered later if needed)
    total_invoices = Invoice.objects.count()
    pending_invoices = Invoice.objects.exclude(status='paid').count()
    
    # Get recent invoices for the dashboard card (latest 10)
    recent_invoices = Invoice.objects.select_related('company__company').order_by('-invoice_date', '-id')[:10]
    
    # Get recent purchase orders (role-based filtering)
    recent_purchase_orders = PurchaseOrder.objects.select_related('company__company').filter(**po_filter).order_by('-order_date', '-id')[:10]
    total_purchase_orders = PurchaseOrder.objects.filter(**po_filter).count()
    
    # Get recent quotations (role-based filtering) - using InquiryHandler for running projects
    if user and user.userprofile.get_roles_list() == 'sales':
        # Sales user - only their assigned inquiries, excluding closed/lost projects
        recent_quotations = InquiryHandler.objects.select_related('company__company').filter(
            sales=user
        ).exclude(
            status__in=['Project Closed', 'Lost']
        ).order_by('-date_of_quote')[:10]
        total_quotations = InquiryHandler.objects.filter(
            sales=user
        ).exclude(
            status__in=['Project Closed', 'Lost']
        ).count()
    else:
        # Admin or no user specified - all inquiries, excluding closed/lost projects
        recent_quotations = InquiryHandler.objects.select_related('company__company').exclude(
            status__in=['Project Closed', 'Lost']
        ).order_by('-date_of_quote')[:10]
        total_quotations = InquiryHandler.objects.exclude(
            status__in=['Project Closed', 'Lost']
        ).count()
    
    # Sales-specific inquiry status chart (filtered by user if applicable)
    sales_statuses = ['Enquiry', 'Inputs', 'Quotation', 'Negotiation', 'PO-Confirm', 'Lost']
    sales_status_counts = {}
    for status in sales_statuses:
        filter_params = {'status': status}
        if inquiry_filter:
            filter_params.update(inquiry_filter)
        sales_status_counts[status] = InquiryHandler.objects.filter(**filter_params).count()
    
    sales_chart_data = {
        'categories': json.dumps(sales_statuses),
        'data': json.dumps([sales_status_counts[status] for status in sales_statuses])
    }
    
    # Calculate Total Payment: Sum of (Order Value × Sales Percentage) for user's POs
    
    # Calculate total payment based on sales percentage
    if user and user.userprofile.get_roles_list() == 'sales':
        # For sales users - only their assigned POs
        total_payment_result = PurchaseOrder.objects.filter(
            sales_person=user,
            sales_percentage__isnull=False  # Only include POs with sales percentage
        ).aggregate(
            total_payment=Sum(
                F('order_value') * F('sales_percentage') / 100,
                output_field=DecimalField(max_digits=15, decimal_places=2)
            )
        )
    else:
        # For admin/manager - all POs with sales percentage
        total_payment_result = PurchaseOrder.objects.filter(
            sales_percentage__isnull=False  # Only include POs with sales percentage
        ).aggregate(
            total_payment=Sum(
                F('order_value') * F('sales_percentage') / 100,
                output_field=DecimalField(max_digits=15, decimal_places=2)
            )
        )
    
    total_payment = total_payment_result['total_payment'] or Decimal('0')
    
    return {
        'dashboard_type': 'sales',
        'total_inquiries': total_inquiries,
        'active_inquiries': active_inquiries,
        'quotations_sent': quotations_sent,
        'total_invoices': total_invoices,
        'pending_invoices': pending_invoices,
        'recent_invoices': recent_invoices,
        'recent_purchase_orders': recent_purchase_orders,
        'total_purchase_orders': total_purchase_orders,
        'recent_quotations': recent_quotations,
        'total_quotations': total_quotations,
        'sales_chart_data': sales_chart_data,
        'is_user_specific': is_user_specific,
        'total_payment': format_indian_currency(total_payment),
        'total_payment_raw': total_payment,
    }


def get_project_manager_dashboard_data(user=None):
    """Get project manager-focused dashboard data"""
    from decimal import Decimal
    from django.db.models import Sum, F, DecimalField
    from django.db.models.functions import Coalesce
    
    def format_indian_currency(amount):
        if amount == 0:
            return "0"
        amount_float = float(amount)
        return f"{amount_float:,.0f}"
    
    # Project management metrics and recent additional supplies - filter by project manager if user is provided
    if user and user.userprofile.get_roles_list() == 'project_manager':
        # Filter Additional Supplies by:
        # 1. Invoice must be generated/active (not draft)
        # 2. PO must be assigned to this project manager
        total_additional_supplies = AdditionalSupply.objects.filter(
            invoice__status__in=['sent', 'invoiced', 'paid', 'partial'],  # Only generated invoices
            invoice__purchase_order__project_manager=user  # Only POs assigned to this PM
        ).count()
        
        total_supply_value = AdditionalSupply.objects.filter(
            invoice__status__in=['sent', 'invoiced', 'paid', 'partial'],  # Only generated invoices
            invoice__purchase_order__project_manager=user  # Only POs assigned to this PM
        ).aggregate(total=Sum('total_amount'))['total'] or Decimal('0')
        
        # Recent additional supplies for this project manager (with both conditions)
        recent_supplies = AdditionalSupply.objects.select_related(
            'invoice__purchase_order', 
            'invoice__company__company'
        ).filter(
            invoice__status__in=['sent', 'invoiced', 'paid', 'partial'],  # Only generated invoices
            invoice__purchase_order__project_manager=user  # Only POs assigned to this PM
        ).order_by('-created_at')[:5]
    else:
        # Show all Additional Supplies from generated invoices (for admin/manager or fallback)
        total_additional_supplies = AdditionalSupply.objects.filter(
            invoice__status__in=['sent', 'invoiced', 'paid', 'partial']  # Only generated invoices
        ).count()
        total_supply_value = AdditionalSupply.objects.filter(
            invoice__status__in=['sent', 'invoiced', 'paid', 'partial']  # Only generated invoices
        ).aggregate(total=Sum('total_amount'))['total'] or Decimal('0')
        
        # Recent additional supplies from generated invoices
        recent_supplies = AdditionalSupply.objects.select_related(
            'invoice__purchase_order', 
            'invoice__company__company'
        ).filter(
            invoice__status__in=['sent', 'invoiced', 'paid', 'partial']  # Only generated invoices
        ).order_by('-created_at')[:5]
    
    # Project status from inquiries
    project_statuses = ['Design', 'Design Review', 'Manufacturing', 'Stage-Inspection', 'Dispatch']
    project_counts = {}
    for status in project_statuses:
        project_counts[status] = InquiryHandler.objects.filter(status=status).count()
    
    # Get recent invoices - filter by project manager if user is provided
    if user and user.userprofile.get_roles_list() == 'project_manager':
        # Filter invoices by:
        # 1. Invoice must be generated/active (not draft)
        # 2. PO must be assigned to this project manager
        recent_invoices = Invoice.objects.select_related('company__company', 'purchase_order').filter(
            status__in=['sent', 'invoiced', 'paid', 'partial'],  # Only generated invoices
            purchase_order__project_manager=user  # Only POs assigned to this PM
        ).order_by('-invoice_date', '-id')[:10]
        total_invoices = Invoice.objects.filter(
            status__in=['sent', 'invoiced', 'paid', 'partial'],  # Only generated invoices
            purchase_order__project_manager=user  # Only POs assigned to this PM
        ).count()
    else:
        # Show all generated invoices (for admin/manager or fallback)
        recent_invoices = Invoice.objects.select_related('company__company', 'purchase_order').filter(
            status__in=['sent', 'invoiced', 'paid', 'partial']  # Only generated invoices
        ).order_by('-invoice_date', '-id')[:10]
        total_invoices = Invoice.objects.filter(
            status__in=['sent', 'invoiced', 'paid', 'partial']  # Only generated invoices
        ).count()
    
    # Get recent purchase orders - filter by project manager if user is provided
    if user and user.userprofile.get_roles_list() == 'project_manager':
        # Filter POs assigned to this project manager
        recent_purchase_orders = PurchaseOrder.objects.select_related('company__company').filter(
            project_manager=user
        ).order_by('-order_date', '-id')[:10]
        total_purchase_orders = PurchaseOrder.objects.filter(project_manager=user).count()
        is_user_specific = True
    else:
        # Show all POs (for admin/manager or fallback)
        recent_purchase_orders = PurchaseOrder.objects.select_related('company__company').order_by('-order_date', '-id')[:10]
        total_purchase_orders = PurchaseOrder.objects.count()
        is_user_specific = False
    
    # Get recent quotations - filter by project manager if user is provided
    if user and user.userprofile.get_roles_list() == 'project_manager':
        # Get companies that have POs assigned to this project manager
        pm_companies = PurchaseOrder.objects.filter(project_manager=user).values_list('company', flat=True)
        
        # Filter quotations for companies that have POs assigned to this PM
        recent_quotations = InquiryHandler.objects.select_related('company__company').filter(
            company__in=pm_companies
        ).exclude(
            status__in=['Project Closed', 'Lost']
        ).order_by('-date_of_quote')[:10]
        
        total_quotations = InquiryHandler.objects.filter(
            company__in=pm_companies
        ).exclude(
            status__in=['Project Closed', 'Lost']
        ).count()
    else:
        # Show all quotations (for admin/manager or fallback)
        recent_quotations = InquiryHandler.objects.select_related('company__company').exclude(
            status__in=['Project Closed', 'Lost']
        ).order_by('-date_of_quote')[:10]
        total_quotations = InquiryHandler.objects.exclude(
            status__in=['Project Closed', 'Lost']
        ).count()
    
    # Calculate Total Payment: Sum of (Order Value × Project Manager Percentage) for user's POs
    
    # Calculate total payment based on project manager percentage
    if user and user.userprofile.get_roles_list() == 'project_manager':
        # For project manager users - only their assigned POs
        total_payment_result = PurchaseOrder.objects.filter(
            project_manager=user,
            project_manager_percentage__isnull=False  # Only include POs with project manager percentage
        ).aggregate(
            total_payment=Sum(
                F('order_value') * F('project_manager_percentage') / 100,
                output_field=DecimalField(max_digits=15, decimal_places=2)
            )
        )
    else:
        # For admin/manager - all POs with project manager percentage
        total_payment_result = PurchaseOrder.objects.filter(
            project_manager_percentage__isnull=False  # Only include POs with project manager percentage
        ).aggregate(
            total_payment=Sum(
                F('order_value') * F('project_manager_percentage') / 100,
                output_field=DecimalField(max_digits=15, decimal_places=2)
            )
        )
    
    total_payment = total_payment_result['total_payment'] or Decimal('0')
    
    return {
        'dashboard_type': 'project_manager',
        'total_additional_supplies': total_additional_supplies,
        'total_supply_value': format_indian_currency(total_supply_value),
        'project_counts': project_counts,
        'recent_supplies': recent_supplies,
        'recent_invoices': recent_invoices,
        'total_invoices': total_invoices,
        'recent_purchase_orders': recent_purchase_orders,
        'total_purchase_orders': total_purchase_orders,
        'recent_quotations': recent_quotations,
        'total_quotations': total_quotations,
        'is_user_specific': is_user_specific,
        'total_payment': format_indian_currency(total_payment),
        'total_payment_raw': total_payment,
    }



@login_required
def user_management_view(request):
    """User management page with list of users"""
    search_query = request.GET.get('search', '')
    users = User.objects.select_related('userprofile').all()
    
    if search_query:
        users = users.filter(
            Q(username__icontains=search_query) |
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(email__icontains=search_query)
        )
    
    users = users.order_by('-date_joined')
    
    # Pagination
    paginator = Paginator(users, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'total_users': User.objects.count(),
        'active_users': User.objects.filter(is_active=True).count(),
        'sales_users': UserProfile.objects.filter(roles__contains='sales').count(),
        'pm_users': UserProfile.objects.filter(roles__contains='project_manager').count(),
    }
    
    return render(request, 'dashboard/user_management.html', context)

@login_required
def user_create_view(request):
    """Create new user with profile"""
    if request.method == 'POST':
        form = UserManagementForm(request.POST)
        if form.is_valid():
            try:
                # Create User
                name_parts = form.cleaned_data['name'].split(' ', 1)
                first_name = name_parts[0]
                last_name = name_parts[1] if len(name_parts) > 1 else ''
                
                # Generate username from name (not email)
                name_for_username = form.cleaned_data['name'].lower().replace(' ', '')
                # Remove any special characters and keep only alphanumeric
                import re
                username = re.sub(r'[^a-zA-Z0-9]', '', name_for_username)
                
                # Ensure username is not empty and has minimum length
                if not username or len(username) < 3:
                    # Fallback to email if name is too short or invalid
                    username = form.cleaned_data['email'].split('@')[0]
                
                # Check for uniqueness and add counter if needed
                counter = 1
                original_username = username
                while User.objects.filter(username=username).exists():
                    username = f"{original_username}{counter}"
                    counter += 1
                
                # Store the plain password for later display
                plain_password = form.cleaned_data['password']
                
                user = User.objects.create_user(
                    username=username,
                    email=form.cleaned_data['email'],
                    password=plain_password,
                    first_name=first_name,
                    last_name=last_name
                )
                
                # Store the password for display purposes
                password_storage.store_password(username, plain_password)
                
                # Create UserProfile with multiple form permissions
                selected_permissions = form.cleaned_data.get('form_permissions', [])
                selected_role = form.cleaned_data['role']
                phone_number = form.cleaned_data.get('phone_number', '')
                
                profile, created = UserProfile.objects.get_or_create(
                    user=user,
                    defaults={
                        'roles': selected_role,
                        'phone_number': phone_number,
                        'can_access_invoice_generation': 'invoice_generation' in selected_permissions,
                        'can_access_inquiry_handler': 'inquiry_handler' in selected_permissions,
                        'can_access_quotation_generation': 'quotation_generation' in selected_permissions,
                        'can_access_additional_supply': 'additional_supply' in selected_permissions
                    }
                )
                
                messages.success(request, f'User {form.cleaned_data["name"]} created successfully!')
                return redirect('dashboard:user_management')
                
            except Exception as e:
                messages.error(request, f'Error creating user: {str(e)}')
                # If user was created but profile failed, clean up
                if 'user' in locals():
                    user.delete()
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = UserManagementForm()
    
    return render(request, 'dashboard/user_form.html', {
        'form': form,
        'title': 'Add User',
        'action': 'Create'
    })

@login_required
def user_edit_view(request, user_id):
    """Edit existing user"""
    user = get_object_or_404(User, id=user_id)
    profile, created = UserProfile.objects.get_or_create(user=user)
    
    if request.method == 'POST':
        form = UserEditForm(request.POST, user_instance=user)
        if form.is_valid():
            # Update User
            name_parts = form.cleaned_data['name'].split(' ', 1)
            user.first_name = name_parts[0]
            user.last_name = name_parts[1] if len(name_parts) > 1 else ''
            user.email = form.cleaned_data['email']
            if form.cleaned_data['password']:
                new_password = form.cleaned_data['password']
                user.set_password(new_password)
                # Update stored password
                password_storage.store_password(user.username, new_password)
            user.save()
            
            # Update UserProfile with single role and multiple form permissions
            selected_role = form.cleaned_data['role']
            selected_permissions = form.cleaned_data['form_permissions']
            phone_number = form.cleaned_data.get('phone_number', '')
            profile.roles = selected_role
            profile.phone_number = phone_number
            profile.can_access_invoice_generation = 'invoice_generation' in selected_permissions
            profile.can_access_inquiry_handler = 'inquiry_handler' in selected_permissions
            profile.can_access_quotation_generation = 'quotation_generation' in selected_permissions
            profile.can_access_additional_supply = 'additional_supply' in selected_permissions
            profile.save()
            
            messages.success(request, f'User {user.get_full_name()} updated successfully!')
            return redirect('dashboard:user_management')
    else:
        # Pre-populate form with selected permissions (multiple)
        selected_permissions = []
        if profile.can_access_invoice_generation:
            selected_permissions.append('invoice_generation')
        if profile.can_access_inquiry_handler:
            selected_permissions.append('inquiry_handler')
        if profile.can_access_quotation_generation:
            selected_permissions.append('quotation_generation')
        if profile.can_access_additional_supply:
            selected_permissions.append('additional_supply')
            
        initial_data = {
            'name': f"{user.first_name} {user.last_name}".strip(),
            'email': user.email,
            'phone_number': profile.phone_number or '',
            'role': profile.get_roles_list(),
            'form_permissions': selected_permissions,
        }
        form = UserEditForm(initial=initial_data, user_instance=user)
    
    return render(request, 'dashboard/user_form.html', {
        'form': form,
        'title': f'Edit User: {user.get_full_name() or user.username}',
        'action': 'Update',
        'user_obj': user,
        'is_edit': True,
        'current_password': password_storage.get_password(user.username)  # Get stored password
    })

@login_required
def user_delete_view(request, user_id):
    """Delete user"""
    user = get_object_or_404(User, id=user_id)
    
    if request.user.id == user.id:
        messages.error(request, 'You cannot delete your own account!')
        return redirect('dashboard:user_management')
    
    if request.method == 'POST':
        username = user.username
        user.delete()
        messages.success(request, f'User {username} deleted successfully!')
        return redirect('dashboard:user_management')
    
    return render(request, 'dashboard/user_delete.html', {'user_obj': user})

def register_view(request):
    """User registration view"""
    if request.user.is_authenticated:
        return redirect('dashboard:index')
    
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            username = form.cleaned_data.get('username')
            messages.success(request, f'Account created for {username}! You can now log in.')
            return redirect('login')
    else:
        form = CustomUserCreationForm()
    
    return render(request, 'registration/register.html', {'form': form})

def login_view(request):
    """Custom login view to handle redirects"""
    if request.user.is_authenticated:
        return redirect('dashboard:index')
    
    # Django's built-in LoginView will handle the actual login
    # This is just for redirect logic
    return redirect('login')

def custom_logout_view(request):
    """Custom logout view that redirects to dashboard URL as requested"""
    logout(request)
    messages.success(request, 'You have been successfully logged out.')
    return redirect('/')  # Redirect to dashboard URL (/) as requested


# Company Management Forms and Views
class CompanyForm(forms.ModelForm):
    class Meta:
        model = Company
        fields = ['company_name', 'city', 'address']
        widgets = {
            'company_name': forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Enter company name'}),
            'city': forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Enter city'}),
            'address': forms.Textarea(attrs={'class': 'form-control', 'rows': 3, 'placeholder': 'Enter company address'}),
        }
    
    def clean(self):
        cleaned_data = super().clean()
        company_name = cleaned_data.get('company_name')
        city = cleaned_data.get('city')
        
        # Check for unique (company_name, city) combination
        # Same company cannot have the same city twice, but different companies can share cities
        if company_name and city:
            existing = Company.objects.filter(company_name=company_name, city=city)
            if self.instance and self.instance.pk:
                existing = existing.exclude(pk=self.instance.pk)
            if existing.exists():
                raise forms.ValidationError(f"This company already exists for the selected city. Company '{company_name}' is already registered in '{city}'.")
        
        return cleaned_data

@login_required
def company_management_view(request):
    """Company management page with list of companies"""
    search_query = request.GET.get('search', '')
    companies = Company.objects.all()
    
    if search_query:
        companies = companies.filter(
            Q(company_name__icontains=search_query) |
            Q(city__icontains=search_query)
        )
    
    companies = companies.order_by('company_name', 'city')
    
    # Pagination
    paginator = Paginator(companies, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'total_companies': Company.objects.count(),
    }
    
    return render(request, 'dashboard/company_management.html', context)

@login_required
def company_create_view(request):
    """Create new company"""
    if request.method == 'POST':
        form = CompanyForm(request.POST)
        if form.is_valid():
            company = form.save()
            messages.success(request, f'Company {company.company_name} created successfully!')
            return redirect('dashboard:contact_management')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = CompanyForm()
    
    return render(request, 'dashboard/company_form.html', {
        'form': form,
        'title': 'Add New Company',
        'action': 'Create'
    })

@login_required
def company_edit_view(request, company_id):
    """Edit existing company"""
    company = get_object_or_404(Company, id=company_id)
    
    if request.method == 'POST':
        form = CompanyForm(request.POST, instance=company)
        if form.is_valid():
            company = form.save()
            messages.success(request, f'Company {company.company_name} updated successfully!')
            return redirect('dashboard:contact_management')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = CompanyForm(instance=company)
    
    return render(request, 'dashboard/company_form.html', {
        'form': form,
        'title': f'Edit Company: {company.company_name}',
        'action': 'Update',
        'company_obj': company,
        'is_edit': True
    })

@login_required
def company_delete_view(request, company_id):
    """Delete company"""
    company = get_object_or_404(Company, id=company_id)
    
    # Check if company is being used by contacts
    contacts_count = Contact.objects.filter(company=company).count()
    
    if request.method == 'POST':
        # Handle AJAX request
        if request.headers.get('Content-Type') == 'application/json':
            if contacts_count > 0:
                return JsonResponse({
                    'success': False,
                    'message': f'Cannot delete company {company.company_name}. It is being used by {contacts_count} contact(s).'
                })
            
            company_name = company.company_name
            company.delete()
            return JsonResponse({
                'success': True,
                'message': f'Company {company_name} deleted successfully!'
            })
        
        # Handle regular form submission (fallback)
        if contacts_count > 0:
            messages.error(request, f'Cannot delete company {company.company_name}. It is being used by {contacts_count} contact(s).')
            return redirect('dashboard:contact_management')
        
        company_name = company.company_name
        company.delete()
        messages.success(request, f'Company {company_name} deleted successfully!')
        return redirect('dashboard:contact_management')
    
    # GET request - show delete confirmation page (fallback)
    return render(request, 'dashboard/company_delete.html', {'company_obj': company})

# Contact Management Forms and Views
class ContactForm(forms.ModelForm):
    # Company selection dropdown
    company = forms.ModelChoiceField(
        queryset=Company.objects.all(),
        empty_label="Select Company",
        widget=forms.Select(attrs={'class': 'form-select', 'id': 'company-select'}),
        label="Company Name *",
        help_text="Select company from master list"
    )
    
    class Meta:
        model = Contact
        fields = ['contact_name', 'email_1', 'phone_1', 'phone_2', 'company', 'individual_address']
        widgets = {
            'contact_name': forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Enter client name'}),
            'email_1': forms.EmailInput(attrs={'class': 'form-control', 'placeholder': 'Enter email address'}),
            'phone_1': forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Enter phone number'}),
            'phone_2': forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Enter secondary phone number (optional)'}),
            'individual_address': forms.Textarea(attrs={'class': 'form-control', 'rows': 3, 'placeholder': 'Enter personal address'}),
        }
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # Order companies by name and city
        self.fields['company'].queryset = Company.objects.all().order_by('company_name', 'city')
        
        # If editing existing contact, set the company field
        if self.instance and self.instance.pk and self.instance.company:
            self.fields['company'].initial = self.instance.company
    
    def clean_phone_1(self):
        phone = self.cleaned_data.get('phone_1')
        if phone and not phone.replace('+', '').replace('-', '').replace(' ', '').replace('(', '').replace(')', '').isdigit():
            raise forms.ValidationError("Please enter a valid phone number.")
        return phone
    
    def clean_phone_2(self):
        phone = self.cleaned_data.get('phone_2')
        if phone and not phone.replace('+', '').replace('-', '').replace(' ', '').replace('(', '').replace(')', '').isdigit():
            raise forms.ValidationError("Please enter a valid phone number.")
        return phone

@login_required
def contact_management_view(request):
    """Contact management page with list of contacts and companies"""
    search_query = request.GET.get('search', '')
    view_type = request.GET.get('view', 'contacts')  # 'contacts' or 'companies'
    
    if view_type == 'companies':
        # Show companies
        companies = Company.objects.all()
        
        if search_query:
            companies = companies.filter(
                Q(company_name__icontains=search_query) |
                Q(city__icontains=search_query)
            )
        
        companies = companies.order_by('company_name', 'city')
        
        # Pagination
        paginator = Paginator(companies, 10)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)
        
        context = {
            'page_obj': page_obj,
            'search_query': search_query,
            'view_type': view_type,
            'total_companies': Company.objects.count(),
            'total_contacts': Contact.objects.count(),
        }
    else:
        # Show contacts (default)
        contacts = Contact.objects.select_related('company').all()
        
        if search_query:
            contacts = contacts.filter(
                Q(contact_name__icontains=search_query) |
                Q(customer_name__icontains=search_query) |
                Q(email_1__icontains=search_query) |
                Q(email_2__icontains=search_query) |
                Q(phone_1__icontains=search_query) |
                Q(company__company_name__icontains=search_query) |
                Q(location_city__icontains=search_query)
            )
        
        contacts = contacts.order_by('-created_at')
        
        # Pagination
        paginator = Paginator(contacts, 10)
        page_number = request.GET.get('page')
        page_obj = paginator.get_page(page_number)
        
        context = {
            'page_obj': page_obj,
            'search_query': search_query,
            'view_type': view_type,
            'total_contacts': Contact.objects.count(),
            'total_companies': Company.objects.count(),
        }
    
    return render(request, 'dashboard/contact_management.html', context)

@login_required
def contact_create_view(request):
    """Create new contact"""
    if request.method == 'POST':
        form = ContactForm(request.POST)
        if form.is_valid():
            contact = form.save()
            messages.success(request, f'Contact {contact.customer_name} created successfully!')
            return redirect('dashboard:contact_management')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = ContactForm()
    
    return render(request, 'dashboard/contact_form.html', {
        'form': form,
        'title': 'Add New Contact',
        'action': 'Create'
    })

@login_required
def contact_edit_view(request, contact_id):
    """Edit existing contact"""
    contact = get_object_or_404(Contact, id=contact_id)
    
    if request.method == 'POST':
        form = ContactForm(request.POST, instance=contact)
        if form.is_valid():
            contact = form.save()
            messages.success(request, f'Contact {contact.customer_name} updated successfully!')
            return redirect('dashboard:contact_management')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = ContactForm(instance=contact)
    
    return render(request, 'dashboard/contact_form.html', {
        'form': form,
        'title': f'Edit Contact: {contact.customer_name}',
        'action': 'Update',
        'contact_obj': contact,
        'is_edit': True
    })

@login_required
def contact_delete_view(request, contact_id):
    """Delete contact"""
    contact = get_object_or_404(Contact, id=contact_id)
    
    if request.method == 'POST':
        customer_name = contact.customer_name
        contact.delete()
        messages.success(request, f'Contact {customer_name} deleted successfully!')
        return redirect('dashboard:contact_management')
    
    return render(request, 'dashboard/contact_delete.html', {'contact_obj': contact})


# Purchase Order Management Forms and Views
class PurchaseOrderForm(forms.ModelForm):
    # Add customer_name as a text input for autocomplete
    customer_name_select = forms.CharField(
        max_length=200,
        widget=forms.TextInput(attrs={
            'class': 'form-control', 
            'id': 'customer-select',
            'placeholder': 'Type to search customer name...',
            'autocomplete': 'off'
        }),
        label="Customer Name *",
        help_text="Start typing to search for customers"
    )
    
    # Hidden field to store the selected customer ID
    selected_customer_id = forms.CharField(
        widget=forms.HiddenInput(),
        required=False
    )
    
    class Meta:
        model = PurchaseOrder
        fields = ['po_number', 'order_date', 'customer_name_select', 'selected_customer_id', 'order_value', 'days_to_mfg', 'remarks', 'payment_terms', 'sales_person', 'sales_percentage', 'project_manager', 'project_manager_percentage']
        widgets = {
            'po_number': forms.TextInput(attrs={'class': 'form-control', 'placeholder': 'Enter PO Number'}),
            'order_date': forms.DateInput(attrs={'class': 'form-control', 'type': 'date'}),
            'order_value': forms.NumberInput(attrs={'class': 'form-control', 'placeholder': 'Enter order value', 'step': '0.01'}),
            'days_to_mfg': forms.NumberInput(attrs={'class': 'form-control', 'placeholder': 'Enter manufacturing days'}),
            'remarks': forms.Textarea(attrs={'class': 'form-control', 'rows': 3, 'placeholder': 'Enter remarks (optional)'}),
            'payment_terms': forms.NumberInput(attrs={'class': 'form-control', 'placeholder': 'Enter days (e.g., 45)', 'min': '1'}),
            'sales_person': forms.Select(attrs={'class': 'form-select'}),
            'sales_percentage': forms.NumberInput(attrs={'class': 'form-control', 'placeholder': 'Enter percentage (e.g., 5.50)', 'step': '0.01', 'min': '0', 'max': '100'}),
            'project_manager': forms.Select(attrs={'class': 'form-select'}),
            'project_manager_percentage': forms.NumberInput(attrs={'class': 'form-control', 'placeholder': 'Enter percentage (e.g., 3.25)', 'step': '0.01', 'min': '0', 'max': '100'}),
        }
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        
        # Filter sales persons (users with sales role)
        sales_users = User.objects.filter(userprofile__roles__contains='sales').distinct()
        self.fields['sales_person'].queryset = sales_users
        self.fields['sales_person'].empty_label = "Select Sales Person"
        
        # Filter project managers (users with project_manager role)
        pm_users = User.objects.filter(userprofile__roles__contains='project_manager').distinct()
        self.fields['project_manager'].queryset = pm_users
        self.fields['project_manager'].empty_label = "Select Project Manager"
        
        # If editing existing order, set the customer fields
        if self.instance and self.instance.pk and self.instance.company:
            self.fields['customer_name_select'].initial = self.instance.company.customer_name
            self.fields['selected_customer_id'].initial = self.instance.company.id
            
        # Debug: Print payment terms value when editing
        if self.instance and self.instance.pk:
            print(f"DEBUG: Editing PO {self.instance.po_number}, Payment Terms: {self.instance.payment_terms}")
            
            # Ensure payment_terms field gets the correct initial value
            if self.instance.payment_terms is not None:
                self.fields['payment_terms'].initial = self.instance.payment_terms
    
    def clean(self):
        cleaned_data = super().clean()
        selected_customer_id = cleaned_data.get('selected_customer_id')
        customer_name_select = cleaned_data.get('customer_name_select')
        
        # Validate that a customer was selected
        if not selected_customer_id and customer_name_select:
            raise forms.ValidationError("Please select a valid customer from the dropdown.")
        
        return cleaned_data
    
    def save(self, commit=True):
        instance = super().save(commit=False)
        
        # Set company and customer_name based on selected customer ID
        selected_customer_id = self.cleaned_data.get('selected_customer_id')
        if selected_customer_id:
            try:
                selected_contact = Contact.objects.get(id=selected_customer_id)
                instance.company = selected_contact
                instance.customer_name = selected_contact.customer_name
            except Contact.DoesNotExist:
                pass
        
        if commit:
            instance.save()
        return instance

@login_required
def purchase_order_management_view(request):
    """Purchase Order management page with list of orders"""
    search_query = request.GET.get('search', '')
    orders = PurchaseOrder.objects.select_related('company', 'sales_person', 'project_manager').all()
    
    if search_query:
        orders = orders.filter(
            Q(po_number__icontains=search_query) |
            Q(company__company__icontains=search_query) |
            Q(customer_name__icontains=search_query) |
            Q(sales_person__username__icontains=search_query) |
            Q(project_manager__username__icontains=search_query)
        )
    
    orders = orders.order_by('-created_at')
    
    # Pagination
    paginator = Paginator(orders, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Statistics
    total_orders = PurchaseOrder.objects.count()
    total_value = PurchaseOrder.objects.aggregate(total=models.Sum('order_value'))['total'] or 0
    overdue_orders = PurchaseOrder.objects.filter(due_days__lt=0).count()
    due_today = PurchaseOrder.objects.filter(due_days=0).count()
    
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'total_orders': total_orders,
        'total_value': total_value,
        'overdue_orders': overdue_orders,
        'due_today': due_today,
    }
    
    return render(request, 'dashboard/purchase_order_management.html', context)

@login_required
def purchase_order_create_view(request):
    """Create new purchase order"""
    if request.method == 'POST':
        form = PurchaseOrderForm(request.POST)
        if form.is_valid():
            order = form.save()
            
            # Handle items data
            items_data_json = request.POST.get('items_data')
            print(f"DEBUG: Received items_data: {items_data_json}")
            
            if items_data_json:
                try:
                    items_data = json.loads(items_data_json)
                    print(f"DEBUG: Parsed items_data: {items_data}")
                    
                    total_amount = 0
                    
                    for item_data in items_data:
                        if item_data.get('item_name') and item_data.get('quantity') and item_data.get('price'):
                            # Extract material code from item name if it's in [CODE] format
                            item_name = item_data['item_name']
                            material_code = ''
                            
                            if item_name.startswith('[') and ']' in item_name:
                                parts = item_name.split(']', 1)
                                material_code = parts[0][1:]  # Remove the opening bracket
                                item_name = parts[1].strip() if len(parts) > 1 else item_name
                            
                            item = PurchaseOrderItem.objects.create(
                                purchase_order=order,
                                material_code=material_code,
                                item_name=item_name,
                                quantity=float(item_data['quantity']),
                                price=float(item_data['price'])
                            )
                            total_amount += float(item.amount)
                            print(f"DEBUG: Created item: {item}")
                    
                    # Update order value with calculated total
                    if total_amount > 0:
                        order.order_value = total_amount
                        order.save()
                        print(f"DEBUG: Updated order value to: {total_amount}")
                        
                except (json.JSONDecodeError, ValueError) as e:
                    print(f"DEBUG: Error processing items: {str(e)}")
                    messages.warning(request, f'Items data could not be processed: {str(e)}')
            else:
                print("DEBUG: No items_data received in POST")
            
            messages.success(request, f'Purchase Order {order.po_number} created successfully!')
            return redirect('dashboard:purchase_order_management')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = PurchaseOrderForm()
    
    return render(request, 'dashboard/purchase_order_form.html', {
        'form': form,
        'title': 'Create Order',
        'action': 'Create'
    })

@login_required
def purchase_order_edit_view(request, order_id):
    """Edit existing purchase order"""
    order = get_object_or_404(PurchaseOrder, id=order_id)
    
    if request.method == 'POST':
        form = PurchaseOrderForm(request.POST, instance=order)
        if form.is_valid():
            order = form.save()
            
            # Handle items data
            items_data_json = request.POST.get('items_data')
            print(f"DEBUG EDIT: Received items_data: {items_data_json}")
            
            if items_data_json:
                try:
                    items_data = json.loads(items_data_json)
                    print(f"DEBUG EDIT: Parsed items_data: {items_data}")
                    
                    # Clear existing items
                    order.items.all().delete()
                    print(f"DEBUG EDIT: Cleared existing items")
                    
                    total_amount = 0
                    for item_data in items_data:
                        if item_data.get('item_name') and item_data.get('quantity') and item_data.get('price'):
                            # Extract material code from item name if it's in [CODE] format
                            item_name = item_data['item_name']
                            material_code = ''
                            
                            if item_name.startswith('[') and ']' in item_name:
                                parts = item_name.split(']', 1)
                                material_code = parts[0][1:]  # Remove the opening bracket
                                item_name = parts[1].strip() if len(parts) > 1 else item_name
                            
                            item = PurchaseOrderItem.objects.create(
                                purchase_order=order,
                                material_code=material_code,
                                item_name=item_name,
                                quantity=float(item_data['quantity']),
                                price=float(item_data['price'])
                            )
                            total_amount += float(item.amount)
                            print(f"DEBUG EDIT: Created item: {item}")
                    
                    # Update order value with calculated total
                    if total_amount > 0:
                        order.order_value = total_amount
                        order.save()
                        print(f"DEBUG EDIT: Updated order value to: {total_amount}")
                        
                except (json.JSONDecodeError, ValueError) as e:
                    print(f"DEBUG EDIT: Error processing items: {str(e)}")
                    messages.warning(request, f'Items data could not be processed: {str(e)}')
            else:
                print("DEBUG EDIT: No items_data received in POST")
            
            messages.success(request, f'Purchase Order {order.po_number} updated successfully!')
            return redirect('dashboard:purchase_order_management')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = PurchaseOrderForm(instance=order)
    
    return render(request, 'dashboard/purchase_order_form.html', {
        'form': form,
        'title': f'Edit Purchase Order: {order.po_number}',
        'action': 'Update',
        'order_obj': order,
        'is_edit': True
    })

@login_required
def purchase_order_delete_view(request, order_id):
    """Delete purchase order"""
    order = get_object_or_404(PurchaseOrder, id=order_id)
    
    if request.method == 'POST':
        po_number = order.po_number
        order.delete()
        messages.success(request, f'Purchase Order {po_number} deleted successfully!')
        return redirect('dashboard:purchase_order_management')
    
    return render(request, 'dashboard/purchase_order_delete.html', {'order_obj': order})

@login_required
def get_company_data_ajax(request):
    """Centralized AJAX endpoint to get company data by company ID"""
    company_id = request.GET.get('company_id')
    if not company_id:
        return JsonResponse({'success': False, 'error': 'No company ID provided'})
    
    try:
        company = Company.objects.get(id=company_id)
        return JsonResponse({
            'success': True,
            'company_id': company.id,
            'company_name': company.company_name,
            'primary_city': company.city,
            'cities_display': company.get_cities_display(),
            'location_city': company.get_primary_city(),
            'addresses': company.get_addresses_list()
        })
    except Company.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Company not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Error fetching company data: {str(e)}'})

@login_required
def get_contact_data_ajax(request):
    """Centralized AJAX endpoint to get contact data by contact ID"""
    contact_id = request.GET.get('contact_id')
    if not contact_id:
        return JsonResponse({'success': False, 'error': 'No contact ID provided'})
    
    try:
        contact = Contact.objects.select_related('company').get(id=contact_id)
        return JsonResponse({
            'success': True,
            'contact_id': contact.id,
            'contact_name': contact.contact_name,
            'customer_name': contact.customer_name,
            'email_1': contact.email_1,
            'email_2': contact.email_2 or '',
            'phone_1': contact.phone_1,
            'phone_2': contact.phone_2 or '',
            'company_id': contact.company.id if contact.company else None,
            'company_name': contact.company.company_name if contact.company else '',
            'location_city': contact.location_city,
            'individual_address': contact.individual_address
        })
    except Contact.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Contact not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Error fetching contact data: {str(e)}'})

@login_required
def export_purchase_orders_excel(request):
    """Export purchase orders to Excel file"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from django.http import HttpResponse
    from datetime import datetime
    
    # Get the same filtered queryset as the management view
    search_query = request.GET.get('search', '')
    orders = PurchaseOrder.objects.select_related('company', 'sales_person', 'project_manager').all()
    
    if search_query:
        orders = orders.filter(
            Q(po_number__icontains=search_query) |
            Q(company__company__icontains=search_query) |
            Q(customer_name__icontains=search_query) |
            Q(sales_person__username__icontains=search_query) |
            Q(project_manager__username__icontains=search_query)
        )
    
    orders = orders.order_by('-created_at')
    
    # Create workbook and worksheet
    wb = Workbook()
    ws = wb.active
    ws.title = "Purchase Orders"
    
    # Define headers
    headers = [
        'PO Number',
        'Company', 
        'Customer',
        'Order Date',
        'Order Value',
        'Delivery Date',
        'Status',
        'Payment Terms',
        'Sales Person',
        'Sales %',
        'Project Manager',
        'PM %'
    ]
    
    # Style for headers
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    
    # Add headers
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
    
    # Add data rows
    for row_num, order in enumerate(orders, 2):
        
        # Company
        ws.cell(row=row_num, column=2, value=order.company.company if order.company else "")
        
        # PO Number
        ws.cell(row=row_num, column=1, value=order.po_number)
        
        # Customer
        ws.cell(row=row_num, column=3, value=order.customer_name)
        
        # Order Date (formatted as DD-MM-YYYY)
        order_date_formatted = order.order_date.strftime('%d-%m-%Y') if order.order_date else ""
        ws.cell(row=row_num, column=4, value=order_date_formatted)
        
        # Order Value (numeric)
        ws.cell(row=row_num, column=5, value=float(order.order_value) if order.order_value else 0)
        
        # Delivery Date (formatted as DD-MM-YYYY)
        delivery_date_formatted = order.delivery_date.strftime('%d-%m-%Y') if order.delivery_date else ""
        ws.cell(row=row_num, column=6, value=delivery_date_formatted)
        
        # Status
        status = order.get_status()
        if order.due_days is not None:
            if order.due_days > 0:
                status = f"{order.due_days} days left"
            elif order.due_days == 0:
                status = "Due today"
            else:
                status = f"{abs(order.due_days)} days overdue"
        ws.cell(row=row_num, column=7, value=status)
        
        # Payment Terms
        payment_terms = f"{order.payment_terms} days" if order.payment_terms else ""
        ws.cell(row=row_num, column=8, value=payment_terms)
        
        # Sales Person
        sales_person_name = ""
        if order.sales_person:
            sales_person_name = order.sales_person.get_full_name() or order.sales_person.username
        ws.cell(row=row_num, column=9, value=sales_person_name)
        
        # Sales Percentage
        sales_percentage = float(order.sales_percentage) if order.sales_percentage else ""
        ws.cell(row=row_num, column=10, value=sales_percentage)
        
        # Project Manager
        pm_name = ""
        if order.project_manager:
            pm_name = order.project_manager.get_full_name() or order.project_manager.username
        ws.cell(row=row_num, column=11, value=pm_name)
        
        # PM Percentage
        pm_percentage = float(order.project_manager_percentage) if order.project_manager_percentage else ""
        ws.cell(row=row_num, column=12, value=pm_percentage)
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)  # Cap at 50 characters
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Create HTTP response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename="purchase_orders.xlsx"'
    
    # Save workbook to response
    wb.save(response)
    
    return response


# Invoice Management Forms and Views
class InvoiceForm(forms.ModelForm):
    # Customer selection autocomplete - primary field for selection
    customer_select = forms.CharField(
        max_length=200,
        widget=forms.TextInput(attrs={
            'class': 'form-control', 
            'id': 'customer-select',
            'placeholder': 'Type to search customer name...',
            'autocomplete': 'off'
        }),
        label="Customer *",
        help_text="Start typing to search for customers"
    )
    
    # Hidden field to store the selected customer ID
    selected_customer_id = forms.CharField(
        widget=forms.HiddenInput(),
        required=False
    )
    
    # Purchase Order selection - will be filtered based on customer selection
    purchase_order = forms.ModelChoiceField(
        queryset=PurchaseOrder.objects.all(),  # Show all initially, will be filtered by JavaScript
        empty_label="Select Purchase Order",
        widget=forms.Select(attrs={'class': 'form-select', 'id': 'purchase-order-select'}),
        label="Purchase Order *"
    )
    
    class Meta:
        model = Invoice
        fields = ['invoice_date', 'customer_select', 'selected_customer_id', 'purchase_order', 'grn_date', 'status', 'remarks']
        # Exclude auto-generated/calculated fields: invoice_number, customer_name, order_value, payment_due_date, due_days, client (auto-fetched)
        widgets = {
            'invoice_date': forms.DateInput(attrs={
                'class': 'form-control', 
                'type': 'date'
            }),
            'grn_date': forms.DateInput(attrs={
                'class': 'form-control', 
                'type': 'date'
            }),
            'status': forms.Select(attrs={
                'class': 'form-select',
                'id': 'payment-status-select'
            }),
            'remarks': forms.Textarea(attrs={
                'class': 'form-control', 
                'rows': 3, 
                'placeholder': 'Enter remarks (optional)'
            }),
        }
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        
        # Limit status choices to only Paid and Partial for payment status
        self.fields['status'].choices = [
            ('paid', 'Paid'),
            ('partial', 'Partial'),
        ]
        self.fields['status'].label = "Payment Status"
        self.fields['status'].help_text = "Select payment status for this invoice"
        
        # If editing existing invoice, set the customer fields
        if self.instance and self.instance.pk and self.instance.company:
            self.fields['customer_select'].initial = self.instance.company.customer_name
            self.fields['selected_customer_id'].initial = self.instance.company.id
            
            # For editing: include current PO even if it has invoice, but exclude others with invoices
            if self.instance.purchase_order:
                self.fields['purchase_order'].queryset = PurchaseOrder.objects.filter(
                    company=self.instance.company
                ).filter(
                    Q(id=self.instance.purchase_order.id) |  # Include current PO
                    ~Q(invoice__isnull=False)  # Exclude POs that have invoices
                ).order_by('-created_at')
                self.fields['purchase_order'].initial = self.instance.purchase_order
            else:
                # If editing but no PO selected, exclude all POs with invoices
                self.fields['purchase_order'].queryset = PurchaseOrder.objects.filter(
                    company=self.instance.company,
                    invoice__isnull=True  # Only POs without invoices
                ).order_by('-created_at')
        else:
            # For new forms, show only purchase orders that don't have invoices yet
            self.fields['purchase_order'].queryset = PurchaseOrder.objects.filter(
                invoice__isnull=True  # Only POs without invoices
            ).order_by('-created_at')
            
            # Set current date as default for new invoices
            from datetime import date
            self.fields['invoice_date'].initial = date.today()
    
    def clean(self):
        cleaned_data = super().clean()
        selected_customer_id = cleaned_data.get('selected_customer_id')
        customer_name_select = cleaned_data.get('customer_select')
        purchase_order = cleaned_data.get('purchase_order')
        
        # Validate that a customer was selected
        if not selected_customer_id and customer_name_select:
            raise forms.ValidationError("Please select a valid customer from the dropdown.")
        
        # Validate that the selected purchase order belongs to the selected customer
        if selected_customer_id and purchase_order:
            try:
                selected_contact = Contact.objects.get(id=selected_customer_id)
                if purchase_order.company != selected_contact:
                    raise forms.ValidationError("The selected purchase order does not belong to the selected customer.")
            except Contact.DoesNotExist:
                raise forms.ValidationError("Selected customer not found.")
        
        return cleaned_data
    
    def save(self, commit=True):
        instance = super().save(commit=False)
        
        # Set company and customer_name based on selected customer ID
        selected_customer_id = self.cleaned_data.get('selected_customer_id')
        if selected_customer_id:
            try:
                selected_contact = Contact.objects.get(id=selected_customer_id)
                instance.company = selected_contact
                instance.customer_name = selected_contact.customer_name
            except Contact.DoesNotExist:
                pass
        
        if commit:
            instance.save()
        return instance

@login_required
def invoice_management_view(request):
    """Invoice management page with list of invoices"""
    search_query = request.GET.get('search', '')
    invoices = Invoice.objects.select_related('company', 'purchase_order').all()
    
    if search_query:
        invoices = invoices.filter(
            Q(invoice_number__icontains=search_query) |
            Q(company__company__icontains=search_query) |
            Q(customer_name__icontains=search_query) |
            Q(purchase_order__po_number__icontains=search_query)
        )
    
    invoices = invoices.order_by('-created_at')
    
    # Pagination
    paginator = Paginator(invoices, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Statistics
    total_invoices = Invoice.objects.count()
    total_value = Invoice.objects.aggregate(total=models.Sum('order_value'))['total'] or 0
    paid_total_value = Invoice.objects.filter(status='paid').aggregate(total=models.Sum('order_value'))['total'] or 0
    partial_total_value = Invoice.objects.filter(status='partial').aggregate(total=models.Sum('order_value'))['total'] or 0
    overdue_invoices = Invoice.objects.filter(due_days__lt=0).count()
    due_today = Invoice.objects.filter(due_days=0).count()
    paid_invoices_count = Invoice.objects.filter(status='paid').count()
    partial_invoices_count = Invoice.objects.filter(status='partial').count()
    
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'total_invoices': total_invoices,
        'total_value': total_value,
        'paid_total_value': paid_total_value,
        'partial_total_value': partial_total_value,
        'overdue_invoices': overdue_invoices,
        'due_today': due_today,
        'paid_invoices_count': paid_invoices_count,
        'partial_invoices_count': partial_invoices_count,
    }
    
    return render(request, 'dashboard/invoice_management.html', context)

@login_required
def invoice_create_view(request):
    """Create Invoice"""
    if request.method == 'POST':
        form = InvoiceForm(request.POST)
        
        # Update purchase order queryset based on selected company
        if 'company' in request.POST and request.POST['company']:
            try:
                company_id = int(request.POST['company'])
                company = Contact.objects.get(id=company_id)
                form.fields['purchase_order'].queryset = PurchaseOrder.objects.filter(
                    company=company
                ).order_by('-created_at')
            except (ValueError, Contact.DoesNotExist):
                pass
        
        if form.is_valid():
            invoice = form.save()
            messages.success(request, f'Invoice {invoice.invoice_number} created successfully!')
            return redirect('dashboard:invoice_management')
        else:
            # Add detailed error messages for debugging
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = InvoiceForm()
    
    return render(request, 'dashboard/invoice_form.html', {
        'form': form,
        'title': 'Create Invoice',
        'action': 'Create'
    })

@login_required
def invoice_edit_view(request, invoice_id):
    """Edit existing invoice"""
    invoice = get_object_or_404(Invoice, id=invoice_id)
    
    if request.method == 'POST':
        form = InvoiceForm(request.POST, instance=invoice)
        
        # Update purchase order queryset based on selected company
        if 'company' in request.POST and request.POST['company']:
            try:
                company_id = int(request.POST['company'])
                company = Contact.objects.get(id=company_id)
                form.fields['purchase_order'].queryset = PurchaseOrder.objects.filter(
                    company=company
                ).order_by('-created_at')
            except (ValueError, Contact.DoesNotExist):
                pass
        
        if form.is_valid():
            invoice = form.save()
            messages.success(request, f'Invoice {invoice.invoice_number} updated successfully!')
            return redirect('dashboard:invoice_management')
        else:
            # Add detailed error messages for debugging
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = InvoiceForm(instance=invoice)
    
    return render(request, 'dashboard/invoice_form.html', {
        'form': form,
        'title': f'Edit Invoice: {invoice.invoice_number}',
        'action': 'Update',
        'invoice_obj': invoice,
        'is_edit': True
    })

@login_required
def invoice_delete_view(request, invoice_id):
    """Delete invoice"""
    invoice = get_object_or_404(Invoice, id=invoice_id)
    
    if request.method == 'POST':
        invoice_number = invoice.invoice_number
        invoice.delete()
        messages.success(request, f'Invoice {invoice_number} deleted successfully!')
        return redirect('dashboard:invoice_management')
    
    return render(request, 'dashboard/invoice_delete.html', {'invoice_obj': invoice})

@login_required
def get_purchase_orders_by_contact_ajax(request):
    """AJAX endpoint to get purchase orders based on selected contact"""
    contact_id = request.GET.get('contact_id')
    invoice_id = request.GET.get('invoice_id')  # For editing case
    
    if not contact_id:
        return JsonResponse({'success': False, 'error': 'No contact ID provided'})
    
    try:
        contact = Contact.objects.select_related('company').get(id=contact_id)
        
        # Base query for purchase orders of this contact
        base_query = PurchaseOrder.objects.filter(company=contact)
        
        if invoice_id:
            # Editing case: include current PO even if it has invoice, exclude others with invoices
            try:
                current_invoice = Invoice.objects.get(id=invoice_id)
                current_po_id = current_invoice.purchase_order.id if current_invoice.purchase_order else None
                
                if current_po_id:
                    # Include current PO + POs without invoices
                    purchase_orders = base_query.filter(
                        Q(id=current_po_id) |  # Include current PO
                        Q(invoice__isnull=True)  # Include POs without invoices
                    ).order_by('-created_at')
                else:
                    # No current PO, just show POs without invoices
                    purchase_orders = base_query.filter(invoice__isnull=True).order_by('-created_at')
            except Invoice.DoesNotExist:
                # Invoice not found, treat as new
                purchase_orders = base_query.filter(invoice__isnull=True).order_by('-created_at')
        else:
            # New invoice case: only POs without invoices
            purchase_orders = base_query.filter(invoice__isnull=True).order_by('-created_at')
        
        po_data = []
        for po in purchase_orders:
            po_data.append({
                'id': po.id,
                'po_number': po.po_number,
                'order_value': str(po.order_value),
                'order_date': po.order_date.strftime('%Y-%m-%d') if po.order_date else '',
                'customer_name': po.customer_name
            })
        
        return JsonResponse({
            'success': True,
            'purchase_orders': po_data,
            'contact_id': contact.id,
            'customer_name': contact.customer_name,
            'company_name': contact.company.company_name if contact.company else ''
        })
    except Contact.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Contact not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Error fetching purchase orders: {str(e)}'})

@login_required
def get_purchase_order_details_ajax(request):
    """AJAX endpoint to get purchase order details"""
    po_id = request.GET.get('po_id')
    if po_id:
        try:
            po = PurchaseOrder.objects.get(id=po_id)
            return JsonResponse({
                'success': True,
                'order_value': str(po.order_value),
                'customer_name': po.customer_name,
                'po_number': po.po_number,
                'payment_terms': po.payment_terms or 15  # Default to 15 days if not set
            })
        except PurchaseOrder.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Purchase Order not found'})
    return JsonResponse({'success': False, 'error': 'No Purchase Order ID provided'})


# Inquiry Handler Management Forms and Views
class InquiryHandlerForm(forms.ModelForm):
    # Customer selection dropdown - primary field for selection
    customer_select = forms.ModelChoiceField(
        queryset=Contact.objects.all(),
        empty_label="Select Customer",
        widget=forms.Select(attrs={'class': 'form-select', 'id': 'customer-select'}),
        label="Customer *",
        help_text="Select customer from database"
    )
    
    # Sales dropdown - replaces BA field (behavior depends on user role)
    sales = forms.ModelChoiceField(
        queryset=User.objects.filter(userprofile__roles__contains='sales'),
        empty_label="Select Sales Person",
        widget=forms.Select(attrs={'class': 'form-select', 'id': 'sales-select'}),
        label="Sales *",
        help_text="Select sales person from database"
    )
    
    # Next Date field - new optional field
    next_date = forms.DateField(
        required=False,
        widget=forms.DateInput(attrs={
            'class': 'form-control',
            'type': 'date'
        }),
        label="Next Date",
        help_text="Enter next date manually (optional)"
    )
    
    class Meta:
        model = InquiryHandler
        fields = ['status', 'customer_select', 'date_of_quote', 'sales', 'next_date', 'remarks']
        # Removed fields: create_id (auto-generated), lead_description (removed), ba (replaced with sales)
        # Auto-generated fields: create_id, opportunity_id, customer_name, quote_no
        widgets = {
            'status': forms.Select(attrs={
                'class': 'form-select',
                'id': 'status-select'
            }),
            'date_of_quote': forms.DateInput(attrs={
                'class': 'form-control',
                'type': 'date',
                'required': 'required'  # Enforce HTML5 required attribute
            }),
            'remarks': forms.Textarea(attrs={
                'class': 'form-control',
                'rows': 3,
                'placeholder': 'Enter remarks (optional)'
            }),
        }
    
    def __init__(self, *args, **kwargs):
        # Extract user from kwargs if provided
        self.user = kwargs.pop('user', None)
        super().__init__(*args, **kwargs)
        
        # Set customer queryset ordered by customer name
        self.fields['customer_select'].queryset = Contact.objects.all().order_by('customer_name')
        
        # Configure sales field based on user role
        if self.user:
            user_role = self.user.userprofile.get_roles_list()
            
            if user_role in ['admin', 'manager']:
                # Admin/Manager: Can select any sales person
                self.fields['sales'].queryset = User.objects.filter(userprofile__roles__contains='sales').order_by('username')
                self.fields['sales'].help_text = "Select sales person from database"
            else:
                # Sales user: Only their own name, not changeable
                self.fields['sales'].queryset = User.objects.filter(userprofile__roles__contains='sales').order_by('username')
                self.fields['sales'].initial = self.user
                self.fields['sales'].widget = forms.HiddenInput()  # Use hidden input instead of disabled
                self.fields['sales'].help_text = f"Assigned to: {self.user.get_full_name() or self.user.username}"
                # Make it required but with only one option
                self.fields['sales'].empty_label = None
        else:
            # Fallback: All sales users
            self.fields['sales'].queryset = User.objects.filter(userprofile__roles__contains='sales').order_by('username')
        
        # Explicitly mark date_of_quote as required
        self.fields['date_of_quote'].required = True
        
        # Set default date to today for new forms
        if not self.instance.pk:
            from datetime import date
            self.fields['date_of_quote'].initial = date.today()
            
            # For new inquiries, set sales person based on user role
            if self.user and self.user.userprofile.get_roles_list() not in ['admin', 'manager']:
                self.fields['sales'].initial = self.user
        
        # If editing existing inquiry, set the customer_select field
        if self.instance and self.instance.pk and self.instance.company:
            self.fields['customer_select'].initial = self.instance.company
    
    def clean_sales(self):
        """Ensure sales field is properly set"""
        sales = self.cleaned_data.get('sales')
        
        # If user is not admin/manager, ensure they can only assign to themselves
        if self.user and self.user.userprofile.get_roles_list() not in ['admin', 'manager']:
            if sales != self.user:
                # Force assignment to current user for non-admin users
                sales = self.user
        
        return sales
    
    def clean_date_of_quote(self):
        """Validate that date_of_quote is provided"""
        date_of_quote = self.cleaned_data.get('date_of_quote')
        if not date_of_quote:
            raise forms.ValidationError("Date of Quote is required and cannot be empty.")
        return date_of_quote
    
    def save(self, commit=True):
        instance = super().save(commit=False)
        
        # Set both company and customer_name based on selected customer
        if self.cleaned_data.get('customer_select'):
            selected_contact = self.cleaned_data['customer_select']
            instance.company = selected_contact  # Set the company relationship
            
            # Ensure customer_name is properly set (fallback to contact_name if customer_name is empty)
            customer_name = selected_contact.customer_name or selected_contact.contact_name
            instance.customer_name = customer_name
        
        if commit:
            instance.save()
        return instance


@login_required
def save_inquiry_items_ajax(request):
    """AJAX endpoint to save inquiry items"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'Invalid request method'})
    
    try:
        import json
        data = json.loads(request.body)
        inquiry_id = data.get('inquiry_id')
        items = data.get('items', [])
        
        if not inquiry_id:
            return JsonResponse({'success': False, 'error': 'No inquiry ID provided'})
        
        inquiry = InquiryHandler.objects.get(id=inquiry_id)
        
        # Clear existing items
        inquiry.items.all().delete()
        
        # Add new items
        total_amount = 0
        for item_data in items:
            if item_data.get('item_name') and item_data.get('quantity') and item_data.get('price'):
                item = InquiryItem.objects.create(
                    inquiry=inquiry,
                    item_name=item_data['item_name'],
                    quantity=float(item_data['quantity']),
                    price=float(item_data['price'])
                )
                total_amount += item.amount
        
        return JsonResponse({
            'success': True,
            'total_amount': str(total_amount),
            'items_count': len(items)
        })
        
    except InquiryHandler.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Inquiry not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Error saving items: {str(e)}'})


@login_required
def get_inquiry_items_ajax(request):
    """AJAX endpoint to get inquiry items"""
    inquiry_id = request.GET.get('inquiry_id')
    if not inquiry_id:
        return JsonResponse({'success': False, 'error': 'No inquiry ID provided'})
    
    try:
        inquiry = InquiryHandler.objects.get(id=inquiry_id)
        items_data = []
        
        for item in inquiry.items.all():
            items_data.append({
                'item_name': item.item_name,
                'quantity': str(item.quantity),
                'price': str(item.price),
                'amount': str(item.amount)
            })
        
        return JsonResponse({
            'success': True,
            'items': items_data
        })
        
    except InquiryHandler.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Inquiry not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Error fetching items: {str(e)}'})

@login_required
def inquiry_handler_management_view(request):
    """Inquiry Handler management page with list of inquiries - filtered by user role and assignments"""
    search_query = request.GET.get('search', '')
    
    # Get user role and permissions
    user_role = request.user.userprofile.get_roles_list()
    
    # Filter inquiries based on user role and permissions
    if user_role in ['admin', 'manager']:
        # Admin and Manager can see all inquiries
        inquiries = InquiryHandler.objects.select_related('company', 'sales').all()
        total_inquiries = InquiryHandler.objects.count()
        active_inquiries = InquiryHandler.objects.exclude(status__in=['Lost', 'Project Closed']).count()
        quotation_stage = InquiryHandler.objects.filter(status='Quotation').count()
        closed_inquiries = InquiryHandler.objects.filter(status='Project Closed').count()
    elif request.user.userprofile.can_access_inquiry_handler:
        # Sales users can only see inquiries assigned to them
        inquiries = InquiryHandler.objects.select_related('company', 'sales').filter(sales=request.user)
        total_inquiries = InquiryHandler.objects.filter(sales=request.user).count()
        active_inquiries = InquiryHandler.objects.filter(sales=request.user).exclude(status__in=['Lost', 'Project Closed']).count()
        quotation_stage = InquiryHandler.objects.filter(sales=request.user, status='Quotation').count()
        closed_inquiries = InquiryHandler.objects.filter(sales=request.user, status='Project Closed').count()
    else:
        # Users without inquiry handler permission see no inquiries
        inquiries = InquiryHandler.objects.none()
        total_inquiries = 0
        active_inquiries = 0
        quotation_stage = 0
        closed_inquiries = 0
    
    # Apply search filter
    if search_query:
        inquiries = inquiries.filter(
            Q(create_id__icontains=search_query) |
            Q(opportunity_id__icontains=search_query) |
            Q(quote_no__icontains=search_query) |
            Q(company__company_name__icontains=search_query) |
            Q(customer_name__icontains=search_query) |
            Q(ba__icontains=search_query) |
            Q(sales__first_name__icontains=search_query) |
            Q(sales__last_name__icontains=search_query) |
            Q(sales__username__icontains=search_query)
        )
    
    inquiries = inquiries.order_by('-created_at')
    
    # Pagination
    paginator = Paginator(inquiries, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'total_inquiries': total_inquiries,
        'active_inquiries': active_inquiries,
        'quotation_stage': quotation_stage,
        'closed_inquiries': closed_inquiries,
        'user_role': user_role,
        'is_admin': user_role in ['admin', 'manager'],
        'can_see_all_inquiries': user_role in ['admin', 'manager'],
    }
    
    return render(request, 'dashboard/inquiry_handler_management.html', context)

@login_required
def inquiry_handler_create_view(request):
    """Create new inquiry"""
    if request.method == 'POST':
        form = InquiryHandlerForm(request.POST, user=request.user)
        
        if form.is_valid():
            inquiry = form.save()
            
            # Handle items data if provided
            items_data = request.POST.get('items_data')
            if items_data:
                try:
                    import json
                    items = json.loads(items_data)
                    
                    # Save each item
                    for item_data in items:
                        if item_data.get('item_name') and item_data.get('quantity') and item_data.get('price'):
                            InquiryItem.objects.create(
                                inquiry=inquiry,
                                item_name=item_data['item_name'],
                                quantity=float(item_data['quantity']),
                                price=float(item_data['price'])
                            )
                except (json.JSONDecodeError, ValueError, KeyError) as e:
                    messages.warning(request, f'Some items could not be saved: {str(e)}')
            
            messages.success(request, f'Inquiry {inquiry.create_id} created successfully!')
            return redirect('dashboard:inquiry_handler_management')
        else:
            # Add detailed error messages for debugging
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = InquiryHandlerForm(user=request.user)
    
    return render(request, 'dashboard/inquiry_handler_form.html', {
        'form': form,
        'title': 'Create New Inquiry',
        'action': 'Create'
    })

@login_required
def inquiry_handler_edit_view(request, inquiry_id):
    """Edit existing inquiry - with user permission checks"""
    inquiry = get_object_or_404(InquiryHandler, id=inquiry_id)
    
    # Check if user has permission to edit this inquiry
    user_role = request.user.userprofile.get_roles_list()
    if user_role not in ['admin', 'manager']:
        # Non-admin users can only edit inquiries assigned to them
        if inquiry.sales != request.user:
            messages.error(request, 'You can only edit inquiries assigned to you.')
            return redirect('dashboard:inquiry_handler_management')
    
    if request.method == 'POST':
        form = InquiryHandlerForm(request.POST, instance=inquiry, user=request.user)
        
        if form.is_valid():
            inquiry = form.save()
            messages.success(request, f'Inquiry {inquiry.create_id} updated successfully!')
            return redirect('dashboard:inquiry_handler_management')
        else:
            # Add detailed error messages for debugging
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = InquiryHandlerForm(instance=inquiry, user=request.user)
    
    return render(request, 'dashboard/inquiry_handler_form.html', {
        'form': form,
        'title': f'Edit Inquiry: {inquiry.create_id}',
        'action': 'Update',
        'inquiry_obj': inquiry,
        'is_edit': True
    })

@login_required
def inquiry_handler_delete_view(request, inquiry_id):
    """Delete inquiry - with user permission checks"""
    inquiry = get_object_or_404(InquiryHandler, id=inquiry_id)
    
    # Check if user has permission to delete this inquiry
    user_role = request.user.userprofile.get_roles_list()
    if user_role not in ['admin', 'manager']:
        # Non-admin users can only delete inquiries assigned to them
        if inquiry.sales != request.user:
            messages.error(request, 'You can only delete inquiries assigned to you.')
            return redirect('dashboard:inquiry_handler_management')
    
    if request.method == 'POST':
        create_id = inquiry.create_id
        inquiry.delete()
        messages.success(request, f'Inquiry {create_id} deleted successfully!')
        return redirect('dashboard:inquiry_handler_management')
    
    return render(request, 'dashboard/inquiry_handler_delete.html', {'inquiry_obj': inquiry})




# Additional Supply Management Forms and Views
class AdditionalSupplyForm(forms.ModelForm):
    # Invoice selection dropdown - primary field for selection
    invoice_select = forms.ModelChoiceField(
        queryset=Invoice.objects.none(),  # Will be set in __init__
        empty_label="Select Invoice",
        widget=forms.Select(attrs={
            'class': 'form-select', 
            'id': 'invoice-select'
        }),
        label="Invoice Number *",
        help_text="Select invoice to auto-fetch details"
    )
    
    class Meta:
        model = AdditionalSupply
        fields = ['invoice_select', 'remarks']
        widgets = {
            'remarks': forms.Textarea(attrs={
                'class': 'form-control',
                'rows': 2,
                'placeholder': 'Enter remarks (optional)'
            }),
        }
    
    def __init__(self, *args, **kwargs):
        # Extract user from kwargs if provided
        user = kwargs.pop('user', None)
        super().__init__(*args, **kwargs)
        
        # Get invoices that don't have additional supplies yet
        invoices_with_supplies = AdditionalSupply.objects.values_list('invoice_id', flat=True).distinct()
        
        # Base queryset for invoices
        base_queryset = Invoice.objects.filter(
            invoice_number__isnull=False
        ).exclude(invoice_number='').select_related('company', 'purchase_order')
        
        # Apply role-based filtering
        if user and hasattr(user, 'userprofile') and user.userprofile.get_roles_list() == 'project_manager':
            # Filter invoices to only show those from POs allocated to this project manager
            base_queryset = base_queryset.filter(
                purchase_order__project_manager=user,
                status__in=['sent', 'invoiced', 'paid', 'partial']  # Only generated invoices
            )
        
        if self.instance and self.instance.pk and self.instance.invoice:
            # If editing, include the current invoice even if it has supplies
            from django.db.models import Q
            self.fields['invoice_select'].queryset = base_queryset.filter(
                Q(id=self.instance.invoice.id) |  # Include current invoice
                ~Q(id__in=invoices_with_supplies)  # Exclude invoices with supplies
            ).order_by('-created_at')
            
            self.fields['invoice_select'].initial = self.instance.invoice
        else:
            # For new forms, exclude all invoices that already have additional supplies
            self.fields['invoice_select'].queryset = base_queryset.exclude(
                id__in=invoices_with_supplies
            ).order_by('-created_at')

@login_required
def additional_supply_management_view(request):
    """Additional Supply management page - shows grouped additional supplies by invoice"""
    search_query = request.GET.get('search', '')
    
    # Get all additional supplies with related data, grouped by invoice
    from django.db.models import Sum, Count
    
    # Group additional supplies by invoice and aggregate data
    invoice_groups = AdditionalSupply.objects.select_related(
        'invoice', 'invoice__company', 'invoice__purchase_order'
    ).values(
        'invoice__id',
        'invoice__invoice_number', 
        'invoice__invoice_date',
        'invoice__company__company__company_name',
        'invoice__customer_name',
        'invoice__purchase_order__po_number'
    ).annotate(
        total_amount=Sum('total_amount'),
        item_count=Count('id'),
        supply_date=models.Min('supply_date')
    ).order_by('-supply_date')
    
    if search_query:
        invoice_groups = invoice_groups.filter(
            models.Q(invoice__invoice_number__icontains=search_query) |
            models.Q(invoice__company__company__company_name__icontains=search_query) |
            models.Q(invoice__customer_name__icontains=search_query) |
            models.Q(invoice__purchase_order__po_number__icontains=search_query)
        )
    
    # Pagination
    paginator = Paginator(invoice_groups, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Statistics
    total_supplies = AdditionalSupply.objects.count()
    total_value = AdditionalSupply.objects.aggregate(total=models.Sum('total_amount'))['total'] or 0
    unique_invoices = AdditionalSupply.objects.values('invoice').distinct().count()
    
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'total_supplies': total_supplies,
        'total_value': total_value,
        'unique_invoices': unique_invoices,
    }
    
    return render(request, 'dashboard/additional_supply_management.html', context)

@login_required
def additional_supply_create_view(request):
    """Create new additional supply with multiple items"""
    if request.method == 'POST':
        form = AdditionalSupplyForm(request.POST, user=request.user)
        
        if form.is_valid():
            # Get the selected invoice
            selected_invoice = form.cleaned_data.get('invoice_select')
            remarks = form.cleaned_data.get('remarks', '')
            
            # Process items from the table
            items_created = 0
            for key, value in request.POST.items():
                if key.startswith('items[') and key.endswith('][description]'):
                    # Extract item index
                    item_index = key.split('[')[1].split(']')[0]
                    
                    # Get item data
                    description = request.POST.get(f'items[{item_index}][description]', '').strip()
                    quantity = request.POST.get(f'items[{item_index}][quantity]', '')
                    unit_price = request.POST.get(f'items[{item_index}][unit_price]', '')
                    
                    # Validate item data
                    if description and quantity and unit_price:
                        try:
                            quantity = float(quantity)
                            unit_price = float(unit_price)
                            
                            # Create additional supply item
                            from datetime import date
                            additional_supply = AdditionalSupply.objects.create(
                                invoice=selected_invoice,
                                supply_date=date.today(),
                                description=description,
                                quantity=quantity,
                                unit_price=unit_price,
                                remarks=remarks
                            )
                            
                            # Create notification for admin dashboard
                            Notification.objects.create(
                                notification_type='additional_supply',
                                title='New Additional Supply Created',
                                message=f'Additional Supply "{description}" created for {selected_invoice.company.company.company_name if selected_invoice.company.company else selected_invoice.company.customer_name}',
                                data={
                                    'po_number': selected_invoice.purchase_order.po_number if selected_invoice.purchase_order else 'N/A',
                                    'customer': selected_invoice.customer_name,
                                    'company': selected_invoice.company.company.company_name if selected_invoice.company.company else selected_invoice.company.customer_name,
                                    'invoice_number': selected_invoice.invoice_number,
                                    'amount': float(additional_supply.total_amount),
                                    'description': description
                                },
                                created_by=request.user
                            )
                            
                            items_created += 1
                        except (ValueError, TypeError):
                            messages.error(request, f'Invalid quantity or unit price for item: {description}')
            
            if items_created > 0:
                messages.success(request, f'{items_created} Additional Supply item(s) for Invoice {selected_invoice.invoice_number} created successfully!')
                return redirect('dashboard:additional_supply_management')
            else:
                messages.error(request, 'No valid items were created. Please check your input.')
        else:
            # Add detailed error messages for debugging
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = AdditionalSupplyForm(user=request.user)
    
    return render(request, 'dashboard/additional_supply_form.html', {
        'form': form,
        'title': 'Create Additional Supply',
        'action': 'Create'
    })

@login_required
def additional_supply_edit_view(request, supply_id):
    """Edit existing additional supply"""
    additional_supply = get_object_or_404(AdditionalSupply, id=supply_id)
    
    if request.method == 'POST':
        form = AdditionalSupplyForm(request.POST, instance=additional_supply, user=request.user)
        
        if form.is_valid():
            # Get the selected invoice
            selected_invoice = form.cleaned_data.get('invoice_select')
            remarks = form.cleaned_data.get('remarks', '')
            
            # Delete existing additional supply record (we'll recreate it with new data)
            additional_supply.delete()
            
            # Process items from the table
            items_created = 0
            for key, value in request.POST.items():
                if key.startswith('items[') and key.endswith('][description]'):
                    # Extract item index
                    item_index = key.split('[')[1].split(']')[0]
                    
                    # Get item data
                    description = request.POST.get(f'items[{item_index}][description]', '').strip()
                    quantity = request.POST.get(f'items[{item_index}][quantity]', '')
                    unit_price = request.POST.get(f'items[{item_index}][unit_price]', '')
                    
                    # Validate item data
                    if description and quantity and unit_price:
                        try:
                            quantity = float(quantity)
                            unit_price = float(unit_price)
                            
                            # Create additional supply item
                            from datetime import date
                            AdditionalSupply.objects.create(
                                invoice=selected_invoice,
                                supply_date=additional_supply.supply_date,  # Keep original supply date
                                description=description,
                                quantity=quantity,
                                unit_price=unit_price,
                                remarks=remarks
                            )
                            items_created += 1
                        except (ValueError, TypeError):
                            messages.error(request, f'Invalid quantity or unit price for item: {description}')
            
            if items_created > 0:
                messages.success(request, f'Additional Supply for Invoice {selected_invoice.invoice_number} updated successfully!')
                return redirect('dashboard:additional_supply_management')
            else:
                messages.error(request, 'No valid items were updated. Please check your input.')
        else:
            # Add detailed error messages for debugging
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = AdditionalSupplyForm(instance=additional_supply, user=request.user)
    
    # Get all additional supplies for the same invoice to populate the table
    existing_supplies = AdditionalSupply.objects.filter(
        invoice=additional_supply.invoice
    ).order_by('id')
    
    return render(request, 'dashboard/additional_supply_form.html', {
        'form': form,
        'title': f'Edit Additional Supply: {additional_supply.invoice.invoice_number}',
        'action': 'Update',
        'supply_obj': additional_supply,
        'existing_supplies': existing_supplies,
        'is_edit': True
    })

@login_required
def additional_supply_delete_view(request, supply_id):
    """Delete additional supply"""
    additional_supply = get_object_or_404(AdditionalSupply, id=supply_id)
    
    if request.method == 'POST':
        invoice_number = additional_supply.invoice.invoice_number
        additional_supply.delete()
        messages.success(request, f'Additional Supply for Invoice {invoice_number} deleted successfully!')
        return redirect('dashboard:additional_supply_management')
    
    return render(request, 'dashboard/additional_supply_delete.html', {'supply_obj': additional_supply})

@login_required
def additional_supply_edit_by_invoice_view(request, invoice_id):
    """Edit all additional supplies for a specific invoice"""
    invoice = get_object_or_404(Invoice, id=invoice_id)
    additional_supplies = AdditionalSupply.objects.filter(invoice=invoice).order_by('id')
    
    if not additional_supplies.exists():
        messages.error(request, f'No additional supplies found for invoice {invoice.invoice_number}')
        return redirect('dashboard:additional_supply_management')
    
    # Use the first supply for the form instance
    first_supply = additional_supplies.first()
    
    if request.method == 'POST':
        form = AdditionalSupplyForm(request.POST, instance=first_supply)
        
        if form.is_valid():
            # Get the selected invoice
            selected_invoice = form.cleaned_data.get('invoice_select')
            remarks = form.cleaned_data.get('remarks', '')
            
            # Delete all existing additional supply records for this invoice
            AdditionalSupply.objects.filter(invoice=invoice).delete()
            
            # Process items from the table
            items_created = 0
            for key, value in request.POST.items():
                if key.startswith('items[') and key.endswith('][description]'):
                    # Extract item index
                    item_index = key.split('[')[1].split(']')[0]
                    
                    # Get item data
                    description = request.POST.get(f'items[{item_index}][description]', '').strip()
                    quantity = request.POST.get(f'items[{item_index}][quantity]', '')
                    unit_price = request.POST.get(f'items[{item_index}][unit_price]', '')
                    
                    # Validate item data
                    if description and quantity and unit_price:
                        try:
                            quantity = float(quantity)
                            unit_price = float(unit_price)
                            
                            # Create additional supply item
                            from datetime import date
                            AdditionalSupply.objects.create(
                                invoice=selected_invoice,
                                supply_date=first_supply.supply_date,  # Keep original supply date
                                description=description,
                                quantity=quantity,
                                unit_price=unit_price,
                                remarks=remarks
                            )
                            items_created += 1
                        except (ValueError, TypeError):
                            messages.error(request, f'Invalid quantity or unit price for item: {description}')
            
            if items_created > 0:
                messages.success(request, f'Additional Supply for Invoice {selected_invoice.invoice_number} updated successfully!')
                return redirect('dashboard:additional_supply_management')
            else:
                messages.error(request, 'No valid items were updated. Please check your input.')
        else:
            # Add detailed error messages for debugging
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = AdditionalSupplyForm(instance=first_supply)
    
    return render(request, 'dashboard/additional_supply_form.html', {
        'form': form,
        'title': f'Edit Additional Supply: {invoice.invoice_number}',
        'action': 'Update',
        'supply_obj': first_supply,
        'existing_supplies': additional_supplies,
        'is_edit': True
    })

@login_required
def additional_supply_delete_by_invoice_view(request, invoice_id):
    """Delete all additional supplies for a specific invoice"""
    invoice = get_object_or_404(Invoice, id=invoice_id)
    additional_supplies = AdditionalSupply.objects.filter(invoice=invoice)
    
    if not additional_supplies.exists():
        messages.error(request, f'No additional supplies found for invoice {invoice.invoice_number}')
        return redirect('dashboard:additional_supply_management')
    
    if request.method == 'POST':
        count = additional_supplies.count()
        additional_supplies.delete()
        messages.success(request, f'All {count} Additional Supply item(s) for Invoice {invoice.invoice_number} deleted successfully!')
        return redirect('dashboard:additional_supply_management')
    
    # Calculate totals for display
    total_amount = sum(supply.total_amount for supply in additional_supplies)
    
    return render(request, 'dashboard/additional_supply_delete_invoice.html', {
        'invoice_obj': invoice,
        'additional_supplies': additional_supplies,
        'total_amount': total_amount,
        'item_count': additional_supplies.count()
    })

@login_required
def get_invoice_details_ajax(request):
    """AJAX endpoint to get invoice details"""
    invoice_id = request.GET.get('invoice_id')
    if not invoice_id:
        return JsonResponse({'success': False, 'error': 'No invoice ID provided'})
    
    try:
        invoice = Invoice.objects.select_related('company', 'purchase_order').get(id=invoice_id)
        
        return JsonResponse({
            'success': True,
            'invoice_number': invoice.invoice_number,
            'po_number': invoice.purchase_order.po_number if invoice.purchase_order else '',
            'company': invoice.company.company.company_name if invoice.company and invoice.company.company else '',
            'customer': invoice.customer_name,
            'order_value': str(invoice.order_value),
            'invoice_date': invoice.invoice_date.strftime('%Y-%m-%d') if invoice.invoice_date else '',
        })
    except Invoice.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Invoice not found'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': f'Error fetching invoice details: {str(e)}'})

# Quotation Generation Views
from django.shortcuts import render
from django.http import HttpResponse

@login_required
@csrf_exempt
def fetch_quotation_data_ajax(request):
    """AJAX endpoint to fetch customer data based on quote number"""
    if request.method == 'POST':
        try:
            import json
            data = json.loads(request.body)
            quote_number = data.get('quote_number', '').strip()
            
            if not quote_number:
                return JsonResponse({'success': False, 'error': 'Quote number is required'})
            
            # Try to find inquiry with matching create_id
            try:
                inquiry = InquiryHandler.objects.select_related('company', 'company__company').get(create_id=quote_number)
                
                # Use company address (simplified - only one address field)
                company_address = ''
                if inquiry.company.company and inquiry.company.company.address:
                    company_address = inquiry.company.company.address
                    address_source = 'company'
                else:
                    # Fallback to individual address if no company address available
                    company_address = inquiry.company.individual_address or ''
                    address_source = 'individual'
                
                # Extract customer data
                customer_data = {
                    'to_person': inquiry.company.contact_name,
                    'firm': inquiry.company.company.company_name if inquiry.company.company else '',
                    'address': company_address,
                    'customer_name': inquiry.customer_name,
                    'email': inquiry.company.email_1 if inquiry.company.email_1 else inquiry.company.email,
                    'phone': inquiry.company.phone_1,
                    'status': inquiry.status,
                    'date_of_quote': inquiry.date_of_quote.strftime('%Y-%m-%d') if inquiry.date_of_quote else '',
                    'address_source': address_source,
                }
                
                return JsonResponse({
                    'success': True,
                    'data': customer_data,
                    'message': f'Customer data loaded for quote {quote_number} (Address: {customer_data["address_source"]})'
                })
                
            except InquiryHandler.DoesNotExist:
                return JsonResponse({
                    'success': False, 
                    'error': f'No inquiry found with quote number: {quote_number}'
                })
                
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'error': 'Invalid JSON data'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': f'Error fetching data: {str(e)}'})
    
    return JsonResponse({'success': False, 'error': 'Invalid request method'})
from django.conf import settings
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from docx.shared import Inches, Mm
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_COLOR_INDEX
from copy import deepcopy

# Fixed image width to insert into table cells (in millimetres). Adjust to 60/70/80 as needed.
IMAGE_WIDTH_MM = 38.1

def seek_and_replace(doc, replacements):
    """Replace text in DOCX while preserving formatting, especially for highlighted text"""
    
    def replace_in_paragraph(paragraph, replacements):
        # Get full text from all runs
        full_text = "".join(run.text for run in paragraph.runs)
        replaced_text = full_text
        
        # Try replacements on full text
        for old, new in replacements.items():
            if new and old in replaced_text:
                replaced_text = replaced_text.replace(old, new)
        
        # If text changed, rebuild the paragraph
        if replaced_text != full_text:
            # Check if any run has highlighting
            has_highlighted = any(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in paragraph.runs)
            
            if has_highlighted:
                # Keep first run's formatting and clear others
                first_run = paragraph.runs[0] if paragraph.runs else None
                if first_run:
                    # Clear all runs
                    for run in paragraph.runs[1:]:
                        run.text = ""
                    # Set new text in first run
                    first_run.text = replaced_text
            else:
                # Non-highlighted text - simple replacement
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = replaced_text
    
    def replace_in_table(table, replacements, is_pricing_table=False):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                # Skip Sr column (column 0) in pricing table to preserve sequential numbers
                if is_pricing_table and cell_idx == 0:
                    continue  # Skip Sr column replacement
                
                # Skip image cells (column 3 in pricing table)
                if is_pricing_table and cell_idx == 3:  # Pricing table image column
                    # Check if this is a fixture row (has serial number)
                    if len(row.cells) > 0 and row.cells[0].text.strip().isdigit():
                        continue  # Skip image cell replacement for fixture rows
                
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)
                for inner_table in cell.tables:
                    replace_in_table(inner_table, replacements)
    
    # Replace in paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    
    # Replace in tables
    for table_idx, table in enumerate(doc.tables):
        is_pricing_table = (table_idx == 1)  # Second table is the pricing table
        replace_in_table(table, replacements, is_pricing_table)
    
    # Replace in headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_paragraph(para, replacements)
        for para in section.footer.paragraphs:
            replace_in_paragraph(para, replacements)
    
    print(f"Text replacements completed for {len(replacements)} items")
    return doc

def _insert_image_in_cell(cell, image_bytes):
    """Insert image into cell, clearing existing content first"""
    if not image_bytes:
        print("No image bytes provided")
        return False
    
    try:
        # Check if Pillow is available for image processing
        try:
            from PIL import Image
            print("Pillow is available for image processing")
        except ImportError:
            print("Warning: Pillow not installed. Image upload feature disabled.")
            return False
        
        print(f"Processing image of size: {len(image_bytes)} bytes")
        
        # Create image stream and validate
        img_stream = BytesIO(image_bytes)
        img_stream.seek(0)
        
        # Validate and process image
        try:
            with Image.open(img_stream) as img:
                print(f"Original image: {img.format}, {img.mode}, {img.size}")
                
                # Convert to RGB if necessary
                if img.mode in ('RGBA', 'LA', 'P'):
                    print(f"Converting image from {img.mode} to RGB")
                    img = img.convert('RGB')
                
                # Resize image if too large (max 800px width)
                if img.width > 800:
                    ratio = 800 / img.width
                    new_height = int(img.height * ratio)
                    img = img.resize((800, new_height), Image.Resampling.LANCZOS)
                    print(f"Resized image to: {img.size}")
                
                # Save as JPEG to ensure compatibility
                processed_stream = BytesIO()
                img.save(processed_stream, format='JPEG', quality=85, optimize=True)
                processed_stream.seek(0)
                print(f"Processed image size: {len(processed_stream.getvalue())} bytes")
                
        except Exception as img_error:
            print(f"Error processing image: {img_error}")
            return False
        
        # Clear cell content
        print("Clearing cell content")
        cell.text = ""
        
        # Remove all existing paragraphs
        for paragraph in cell.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Add new paragraph
        paragraph = cell.add_paragraph()
        run = paragraph.add_run()
        
        # Set paragraph alignment to center
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add image to document with proper sizing
        try:
            processed_stream.seek(0)
            run.add_picture(processed_stream, width=Mm(IMAGE_WIDTH_MM))
            print(f"Successfully added image to cell with width {IMAGE_WIDTH_MM}mm")
        except Exception as docx_error:
            print(f"Error adding image to DOCX: {docx_error}")
            return False
        
        # Mark the cell as having an image
        cell._element.set('image_inserted', 'true')
        print("Image insertion completed successfully")
        return True
        
    except Exception as e:
        print(f"Error inserting image: {e}")
        import traceback
        traceback.print_exc()
        return False

def add_fixture_rows_to_table(table, fixtures_data):
    """Add additional fixture rows dynamically for fixtures beyond the first 2"""
    if len(fixtures_data) <= 2:
        return  # No additional rows needed
    
    print(f"=== ADDING ROWS FOR {len(fixtures_data) - 2} ADDITIONAL FIXTURES ===")
    
    # Find column positions from header
    column_positions = {}
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for col_idx, cell in enumerate(header_row.cells):
            header_text = cell.text.strip().lower()
            if 'sr' in header_text or 'serial' in header_text:
                column_positions['sr'] = col_idx
            elif 'description' in header_text:
                column_positions['description'] = col_idx
            elif 'image' in header_text:
                column_positions['image'] = col_idx
            elif 'qty' in header_text or 'quantity' in header_text:
                column_positions['qty'] = col_idx
            elif 'rate' in header_text:
                column_positions['rate'] = col_idx
            elif 'per' in header_text or 'unit' in header_text:
                column_positions['per'] = col_idx
            elif 'amount' in header_text:
                column_positions['amount'] = col_idx
    
    # Template rows to clone (fixture row and its corresponding words row)
    template_fixture_row = table.rows[1] if len(table.rows) > 1 else None  # First fixture row
    template_words_row = table.rows[2] if len(table.rows) > 2 else None    # First words row
    
    if not template_fixture_row:
        print("No template fixture row found")
        return
    
    # Insert position should be at the end of existing rows
    insert_position = len(table.rows)
    
    for idx in range(2, len(fixtures_data)):  # Start from 3rd fixture
        fixture = fixtures_data[idx]
        
        print(f"Adding fixture {idx + 1}")
        
        # Clone fixture row and insert it at the correct position
        new_fixture_row = _clone_row_at_position(table, template_fixture_row, insert_position)
        insert_position += 1
        
        # Fill in the fixture data using column positions
        cells = new_fixture_row.cells
        
        try:
            # Clear all cells first
            for cell in cells:
                cell.text = ""
            
            # Place Sr number in Sr column
            if 'sr' in column_positions and column_positions['sr'] < len(cells):
                cells[column_positions['sr']].text = str(idx + 1)
                print(f"Set Sr to '{idx + 1}' in column {column_positions['sr']}")
            
            # Place description in Description column
            if 'description' in column_positions and column_positions['description'] < len(cells):
                desc_lines = []
                if fixture.get('name') or fixture.get('desc'):
                    desc_lines.append(f"{fixture.get('name', '')} {fixture.get('desc', '')}".strip())
                desc_lines.append("(as per annexure)")
                if fixture.get('hsn'):
                    desc_lines.append(f"HSN Code: {fixture.get('hsn')}")
                if fixture.get('specifications', '').strip():
                    desc_lines.append(f"Specifications: {fixture.get('specifications')}")
                
                description_text = "\n".join(desc_lines)
                cells[column_positions['description']].text = description_text
                print(f"Set description in column {column_positions['description']}")
            
            # Place image in Image column
            if 'image' in column_positions and column_positions['image'] < len(cells):
                if fixture.get('image'):
                    success = _insert_image_in_cell(cells[column_positions['image']], fixture.get('image'))
                    if not success:
                        cells[column_positions['image']].text = f"[Image: {fixture.get('name', 'Product')}]"
                else:
                    cells[column_positions['image']].text = ""
            
            # Place other data in correct columns
            if 'qty' in column_positions and column_positions['qty'] < len(cells):
                cells[column_positions['qty']].text = fixture.get('qty', '1')
            
            if 'rate' in column_positions and column_positions['rate'] < len(cells):
                cells[column_positions['rate']].text = fixture.get('price', '')
            
            if 'per' in column_positions and column_positions['per'] < len(cells):
                cells[column_positions['per']].text = fixture.get('unit', 'Set')
            
            if 'amount' in column_positions and column_positions['amount'] < len(cells):
                cells[column_positions['amount']].text = fixture.get('total', '')
            
            print(f"Successfully added fixture {idx + 1}")
            
        except Exception as e:
            print(f"Error filling fixture data for fixture {idx + 1}: {e}")
            continue
        
        # Clone words row immediately after the fixture row
        if template_words_row:
            new_words_row = _clone_row_at_position(table, template_words_row, insert_position)
            insert_position += 1
            
            # Set the "In Words" text
            try:
                words_cells = new_words_row.cells
                
                # Clear all cells first
                for cell in words_cells:
                    cell.text = ""
                
                # Clear Sr column
                if 'sr' in column_positions and column_positions['sr'] < len(words_cells):
                    words_cells[column_positions['sr']].text = ""
                
                # Set words text starting from Description column
                words_text = f"In Words: {fixture.get('words', '')}"
                start_col = column_positions.get('description', 1)
                for col_idx in range(start_col, len(words_cells)):
                    words_cells[col_idx].text = words_text
                
                print(f"Set words text for fixture {idx + 1}")
                
            except Exception as e:
                print(f"Error setting words text for fixture {idx + 1}: {e}")
    
    print(f"=== COMPLETED ADDING {len(fixtures_data) - 2} ADDITIONAL FIXTURES ===")

def _clone_row_at_position(table, source_row, insert_position):
    """Clone a complete row with all cell properties and merged cells at specific position"""
    from copy import deepcopy
    
    # Clone the row element completely
    new_row_element = deepcopy(source_row._element)
    
    # Get the table element
    tbl_element = table._element
    
    # Insert the new row at the specified position
    if insert_position >= len(table.rows):
        # Append to end
        tbl_element.append(new_row_element)
    else:
        # Insert at specific position
        existing_row = table.rows[insert_position]._element
        tbl_element.insert(tbl_element.index(existing_row), new_row_element)
    
    # Return the newly created row
    # We need to refresh the table rows to get the correct row object
    return table.rows[insert_position]

def populate_pricing_table_with_fixtures(table, fixtures_data):
    """DYNAMIC column-wise data placement - works for ANY number of fixtures"""
    if not fixtures_data:
        return
    
    print(f"=== DYNAMIC TABLE POPULATION FOR {len(fixtures_data)} FIXTURES ===")
    print(f"Table has {len(table.rows)} rows, {len(table.columns)} columns")
    
    # STEP 1: Find column positions by reading headers
    column_positions = {}
    if len(table.rows) > 0:
        header_row = table.rows[0]
        for col_idx, cell in enumerate(header_row.cells):
            header_text = cell.text.strip().lower()
            print(f"Column {col_idx}: '{header_text}'")
            
            if 'sr' in header_text or 'serial' in header_text:
                column_positions['sr'] = col_idx
            elif 'description' in header_text:
                column_positions['description'] = col_idx
            elif 'image' in header_text:
                column_positions['image'] = col_idx
            elif 'qty' in header_text or 'quantity' in header_text:
                column_positions['qty'] = col_idx
            elif 'rate' in header_text:
                column_positions['rate'] = col_idx
            elif 'per' in header_text or 'unit' in header_text:
                column_positions['per'] = col_idx
            elif 'amount' in header_text:
                column_positions['amount'] = col_idx
    
    print(f"Column positions found: {column_positions}")
    
    # STEP 2: Find ALL fixture rows dynamically (not just first 2)
    fixture_rows = []
    for r_idx in range(1, len(table.rows)):  # Skip header row
        row = table.rows[r_idx]
        row_text = " ".join([cell.text.strip() for cell in row.cells]).lower()
        
        # Check for any fixture pattern
        if "fixture" in row_text and ("holding" in row_text or "pulling" in row_text or "connector" in row_text or "pcba" in row_text):
            # Try to determine fixture number from the row
            fixture_num = None
            if "fixture 1" in row_text or "holding" in row_text:
                fixture_num = 0
            elif "fixture 2" in row_text or "pulling" in row_text:
                fixture_num = 1
            else:
                # For fixtures beyond 2, try to extract number or assign sequentially
                for i in range(len(fixtures_data)):
                    if f"fixture {i+1}" in row_text:
                        fixture_num = i
                        break
                
                # If still not found, assign based on order
                if fixture_num is None:
                    fixture_num = len(fixture_rows)
            
            if fixture_num is not None and fixture_num < len(fixtures_data):
                fixture_rows.append((r_idx, fixture_num))
                print(f"Found Fixture {fixture_num + 1} at row {r_idx}")
    
    print(f"Found {len(fixture_rows)} fixture rows for {len(fixtures_data)} fixtures")
    
    # STEP 3: Process ALL fixtures dynamically
    for row_idx, fixture_idx in fixture_rows:
        if fixture_idx >= len(fixtures_data):
            continue
            
        fixture = fixtures_data[fixture_idx]
        row = table.rows[row_idx]
        cells = row.cells
        
        print(f"=== PLACING DATA FOR FIXTURE {fixture_idx + 1} IN ROW {row_idx} ===")
        
        # CLEAR only the data cells (not image cells) to prevent mixing
        for col_name, col_idx in column_positions.items():
            if col_name != 'image' and col_idx < len(cells):
                cells[col_idx].text = ""
        
        # Place Sr number in Sr column
        if 'sr' in column_positions and column_positions['sr'] < len(cells):
            sr_col = column_positions['sr']
            cells[sr_col].text = str(fixture_idx + 1)
            print(f"Placed Sr '{fixture_idx + 1}' in column {sr_col}")
        
        # Place description in Description column
        if 'description' in column_positions and column_positions['description'] < len(cells):
            desc_col = column_positions['description']
            desc_lines = []
            if fixture.get('name') or fixture.get('desc'):
                desc_lines.append(f"{fixture.get('name', '')} {fixture.get('desc', '')}".strip())
            desc_lines.append("(as per annexure)")
            if fixture.get('hsn'):
                desc_lines.append(f"HSN Code: {fixture.get('hsn')}")
            if fixture.get('specifications', '').strip():
                desc_lines.append(f"Specifications: {fixture.get('specifications')}")
            
            description_text = "\n".join(desc_lines)
            cells[desc_col].text = description_text
            print(f"Placed description in column {desc_col}")
        
        # Place image in Image column
        if 'image' in column_positions and column_positions['image'] < len(cells):
            img_col = column_positions['image']
            if fixture.get('image'):
                success = _insert_image_in_cell(cells[img_col], fixture.get('image'))
                if not success:
                    cells[img_col].text = f"[Image: {fixture.get('name', 'Product')}]"
                print(f"Placed image in column {img_col}")
            else:
                cells[img_col].text = ""
        
        # Place quantity in Qty column
        if 'qty' in column_positions and column_positions['qty'] < len(cells):
            qty_col = column_positions['qty']
            cells[qty_col].text = fixture.get('qty', '1')
            print(f"Placed qty '{fixture.get('qty', '1')}' in column {qty_col}")
        
        # Place rate in Rate column
        if 'rate' in column_positions and column_positions['rate'] < len(cells):
            rate_col = column_positions['rate']
            cells[rate_col].text = fixture.get('price', '')
            print(f"Placed rate '{fixture.get('price', '')}' in column {rate_col}")
        
        # Place unit in Per column
        if 'per' in column_positions and column_positions['per'] < len(cells):
            per_col = column_positions['per']
            cells[per_col].text = fixture.get('unit', 'Set')
            print(f"Placed per '{fixture.get('unit', 'Set')}' in column {per_col}")
        
        # Place amount in Amount column
        if 'amount' in column_positions and column_positions['amount'] < len(cells):
            amount_col = column_positions['amount']
            cells[amount_col].text = fixture.get('total', '')
            print(f"Placed amount '{fixture.get('total', '')}' in column {amount_col}")
    
    # STEP 4: Handle "In Words" rows dynamically
    words_rows = []
    for r_idx in range(1, len(table.rows)):
        row = table.rows[r_idx]
        row_text = " ".join([cell.text.strip() for cell in row.cells]).lower()
        if "in words:" in row_text:
            words_rows.append(r_idx)
    
    print(f"Found {len(words_rows)} words rows")
    
    # Process words rows for ALL fixtures
    for i, row_idx in enumerate(words_rows):
        if i < len(fixtures_data):
            fixture = fixtures_data[i]
            row = table.rows[row_idx]
            cells = row.cells
            
            # Clear Sr column
            if 'sr' in column_positions and column_positions['sr'] < len(cells):
                cells[column_positions['sr']].text = ""
            
            # Place words text starting from Description column
            words_text = f"In Words: {fixture.get('words', '')}"
            start_col = column_positions.get('description', 1)
            for col_idx in range(start_col, len(cells)):
                cells[col_idx].text = words_text
            
            print(f"Placed words for fixture {i + 1} in row {row_idx}")
    
    print("=== DYNAMIC TABLE POPULATION COMPLETE ===")
    
    # VERIFICATION: Show what's in each column for all fixtures
    print("=== VERIFICATION FOR ALL FIXTURES ===")
    fixture_count = 0
    for r_idx in range(min(20, len(table.rows))):  # Check more rows for multiple fixtures
        row = table.rows[r_idx]
        cells = row.cells
        
        # Check if this is a fixture row
        if 'sr' in column_positions and column_positions['sr'] < len(cells):
            sr_text = cells[column_positions['sr']].text.strip()
            if sr_text.isdigit():
                fixture_count += 1
                print(f"Fixture {sr_text} (Row {r_idx}):")
                for col_name, col_idx in column_positions.items():
                    if col_idx < len(cells):
                        cell_text = cells[col_idx].text.strip()[:30]
                        if cell_text:
                            print(f"  {col_name} (col {col_idx}): '{cell_text}'")
    
    print(f"Verified {fixture_count} fixtures in table")

def insert_images_in_inclusion_section(doc, fixtures_data):
    """IMPROVED APPROACH: Keep original Inclusions section, just add images properly with correct font size"""
    try:
        print("=== IMPROVED APPROACH: ADDING IMAGES TO EXISTING INCLUSION SECTION ===")
        print(f"Processing {len(fixtures_data)} fixtures")
        
        # Find the inclusion section paragraph that contains the actual fixture content
        inclusion_content_paragraph = None
        inclusion_para_index = -1
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            para_text = paragraph.text.strip()
            # Look for paragraph that contains fixture content (not just the header)
            if ("1. Product Name:" in para_text or 
                "Product Name:" in para_text and len(para_text) > 50):
                inclusion_content_paragraph = paragraph
                inclusion_para_index = para_idx
                print(f"Found inclusion content at paragraph {para_idx}")
                print(f"Original content: '{para_text[:100]}...'")
                break
        
        if not inclusion_content_paragraph:
            print("Could not find inclusion content paragraph")
            return
        
        # Clear the existing content paragraph and rebuild it with images
        inclusion_content_paragraph.clear()
        
        # Build new inclusion content with images in the SAME paragraph location
        from docx.shared import Mm, Pt
        from io import BytesIO
        
        for idx, fixture in enumerate(fixtures_data, 1):
            if not (fixture.get('name') or fixture.get('desc')):
                continue
                
            print(f"Adding fixture {idx}: {fixture.get('name', 'Unknown')}")
            
            # Add the fixture number and product name with 16pt font
            run = inclusion_content_paragraph.add_run(f"{idx}. Product Name: {fixture.get('name', f'Fixture {idx}')}")
            run.font.size = Pt(16)  # Set font size to 16pt
            inclusion_content_paragraph.add_run().add_break()
            
            # Add image if available - RIGHT AFTER Product Name (smaller size)
            if fixture.get('image'):
                try:
                    img_stream = BytesIO(fixture['image'])
                    img_stream.seek(0)
                    img_run = inclusion_content_paragraph.add_run()
                    img_run.add_picture(img_stream, width=Mm(35))  # Reduced from 50mm to 35mm
                    inclusion_content_paragraph.add_run().add_break()
                    print(f"✅ Added image for fixture {idx} (35mm width)")
                except Exception as img_error:
                    print(f"❌ Failed to add image for fixture {idx}: {img_error}")
                    error_run = inclusion_content_paragraph.add_run("   [Image could not be loaded]")
                    error_run.font.size = Pt(16)
                    inclusion_content_paragraph.add_run().add_break()
            
            # Add description with 16pt font
            if fixture.get('desc'):
                desc_run = inclusion_content_paragraph.add_run(f"   Description: {fixture['desc']}")
                desc_run.font.size = Pt(16)
                inclusion_content_paragraph.add_run().add_break()
            
            # Add standard text with 16pt font
            annexure_run = inclusion_content_paragraph.add_run("   (as per annexure)")
            annexure_run.font.size = Pt(16)
            inclusion_content_paragraph.add_run().add_break()
            
            # Add HSN Code with 16pt font
            if fixture.get('hsn'):
                hsn_run = inclusion_content_paragraph.add_run(f"   HSN Code: {fixture['hsn']}")
                hsn_run.font.size = Pt(16)
                inclusion_content_paragraph.add_run().add_break()
            
            # Add specifications with 16pt font
            if fixture.get('specifications', '').strip():
                spec_run = inclusion_content_paragraph.add_run(f"   Specifications: {fixture['specifications']}")
                spec_run.font.size = Pt(16)
                inclusion_content_paragraph.add_run().add_break()
            
            # Add spacing between fixtures (except for the last one)
            if idx < len([f for f in fixtures_data if f.get('name') or f.get('desc')]):
                inclusion_content_paragraph.add_run().add_break()
        
        print("=== INCLUSION SECTION UPDATED WITH 16PT FONT AND 35MM IMAGES ===")
        
    except Exception as e:
        print(f"❌ Error updating inclusion section: {e}")
        import traceback
        traceback.print_exc()

@login_required
def quotation_view(request):
    """Quotation generator page with draft loading capability"""
    # Check for specific draft ID in URL parameter
    draft_id = request.GET.get('draft_id')
    draft_data = None
    
    if draft_id:
        # Load specific quotation by ID (could be draft or generated)
        try:
            specific_quotation = Quotation.objects.get(
                id=draft_id,
                created_by=request.user
            )
            
            # If this is a generated quotation, find the latest revision for the same base quote
            if specific_quotation.status == 'generated':
                base_quote_no = specific_quotation.quote_number.split('_')[0]
                latest_quotation = Quotation.objects.filter(
                    quote_number__startswith=base_quote_no,
                    created_by=request.user
                ).order_by('-updated_at').first()
                
                if latest_quotation and latest_quotation.id != specific_quotation.id:
                    print(f"Loading latest revision: {latest_quotation.quote_number} instead of {specific_quotation.quote_number}")
                    specific_quotation = latest_quotation
            
            # Get images if it's a draft
            quotation_images = {}
            if specific_quotation.status == 'draft':
                for img in specific_quotation.draft_images.all():
                    quotation_images[img.fixture_index] = {
                        'url': img.image.url,
                        'filename': img.original_filename,
                        'size': img.file_size
                    }
            print(f"Loaded {len(quotation_images)} images for quotation {specific_quotation.id}: {quotation_images}")
            
            draft_data = {
                'quote_no': specific_quotation.quote_number,
                'revision': specific_quotation.revision,
                'date': specific_quotation.quotation_date,
                'to_person': specific_quotation.to_person,
                'firm': specific_quotation.firm,
                'address': specific_quotation.address,
                'payment_terms': specific_quotation.payment_terms,
                'delivery_terms': specific_quotation.delivery_terms,
                'fixtures': specific_quotation.fixtures_data or [],
                'draft_id': specific_quotation.id,
                'images': quotation_images,
                'is_draft': specific_quotation.status == 'draft',  # True only if status is 'draft'
                'is_readonly_revision': specific_quotation.status != 'draft'  # Make revision readonly for all non-draft statuses
            }
            print(f"Loaded quotation data for ID {specific_quotation.id}:")
            print(f"  - Status: {specific_quotation.status}")
            print(f"  - Is Draft: {draft_data['is_draft']}")
            print(f"  - Fixtures count: {len(draft_data['fixtures'])}")
            if draft_data['fixtures']:
                print(f"  - First fixture: {draft_data['fixtures'][0]}")
            print(f"  - Images count: {len(quotation_images)}")
        except Quotation.DoesNotExist:
            pass
    else:
        # Check for latest draft for this user
        latest_draft = Quotation.objects.filter(
            status='draft',
            created_by=request.user
        ).order_by('-updated_at').first()
        
        if latest_draft:
            # Get draft images
            draft_images = {}
            for img in latest_draft.draft_images.all():
                draft_images[img.fixture_index] = {
                    'url': img.image.url,
                    'filename': img.original_filename,
                    'size': img.file_size
                }
            print(f"Loaded {len(draft_images)} images for latest draft {latest_draft.id}: {draft_images}")
            
            draft_data = {
                'quote_no': latest_draft.quote_number,
                'revision': latest_draft.revision,
                'date': latest_draft.quotation_date,
                'to_person': latest_draft.to_person,
                'firm': latest_draft.firm,
                'address': latest_draft.address,
                'payment_terms': latest_draft.payment_terms,
                'delivery_terms': latest_draft.delivery_terms,
                'fixtures': latest_draft.fixtures_data or [],
                'draft_id': latest_draft.id,
                'images': draft_images,
                'is_draft': True,  # This is a draft
                'is_readonly_revision': False  # Drafts can edit revision
            }
    
    if draft_data:
        import json
        draft_data_json = json.dumps(draft_data)
        print(f"=== DRAFT DATA DEBUG ===")
        print(f"Draft data keys: {list(draft_data.keys()) if draft_data else 'None'}")
        print(f"Draft data JSON length: {len(draft_data_json) if draft_data_json else 0}")
        print(f"Is draft: {draft_data.get('is_draft', 'Unknown')}")
    else:
        draft_data_json = None
        print("=== NO DRAFT DATA ===")
    
    return render(request, 'dashboard/quotation_generator.html', {
        'draft_data': draft_data,
        'draft_data_json': draft_data_json
    })

@login_required
def generate_quotation(request):
    """Generate quotation document with proper image handling and specifications"""
    if request.method == 'POST':
        try:
            # Collect form data
            quote_data = {
                'quote_no': request.POST.get('quote_no', 'KEC005JN2025'),
                'revision': request.POST.get('revision', 'Rev A'),
                'date': request.POST.get('date', 'Wednesday, September 24, 2025'),
                'to_person': request.POST.get('to_person', 'Mr. Mohak Dholakia'),
                'firm': request.POST.get('firm', 'Schneider Electric India Private Limited'),
                'address': request.POST.get('address', ''),
                'payment_terms': request.POST.get('payment_terms', 'Payment: 45 Days from delivery (Being an MSME)'),
                'delivery_terms': request.POST.get('delivery_terms', 'Delivery: 2-3 weeks per fixture from Purchase Order'),
            }
            
            # Process fixtures with image handling
            fixtures = []
            fixture_index = 0
            
            while True:
                name_key = f'fixtures[{fixture_index}][name]'
                if name_key not in request.POST:
                    break
                    
                fixture = {
                    'name': request.POST.get(name_key, ''),
                    'desc': request.POST.get(f'fixtures[{fixture_index}][desc]', ''),
                    'hsn': request.POST.get(f'fixtures[{fixture_index}][hsn]', '84790000'),
                    'qty': request.POST.get(f'fixtures[{fixture_index}][qty]', '1'),
                    'unit': request.POST.get(f'fixtures[{fixture_index}][unit]', 'Set'),
                    'price': request.POST.get(f'fixtures[{fixture_index}][price]', ''),
                    'total': request.POST.get(f'fixtures[{fixture_index}][total]', ''),
                    'words': request.POST.get(f'fixtures[{fixture_index}][words]', ''),
                    'specifications': request.POST.get(f'fixtures[{fixture_index}][specifications]', ''),
                    'has_image': False,
                    'image': None,
                }
                
                # Handle image upload for this fixture
                image_key = f'fixtures[{fixture_index}][image]'
                if image_key in request.FILES:
                    uploaded_file = request.FILES[image_key]
                    if uploaded_file and uploaded_file.size > 0:
                        try:
                            image_bytes = uploaded_file.read()
                            fixture['image'] = image_bytes
                            fixture['has_image'] = True
                            print(f"Processed image for fixture {fixture_index + 1}: {len(image_bytes)} bytes")
                        except Exception as img_error:
                            print(f"Error processing image for fixture {fixture_index + 1}: {img_error}")
                            fixture['has_image'] = False
                            fixture['image'] = None
                
                fixtures.append(fixture)
                fixture_index += 1
            
            # Ensure we have at least 2 fixtures for template compatibility
            while len(fixtures) < 2:
                fixtures.append({
                    'name': f'Fixture {len(fixtures) + 1}',
                    'desc': '',
                    'hsn': '84790000',
                    'qty': '1',
                    'unit': 'Set',
                    'price': '',
                    'total': '',
                    'words': '',
                    'specifications': '',
                    'has_image': False,
                    'image': None,
                })
            
            # Build content for tag-based replacements with enhanced inclusion section
            inclusion_items = []
            scope_items = []
            specification_items = []
            
            for idx, fixture in enumerate(fixtures, 1):
                if fixture['name'] or fixture['desc']:
                    # Enhanced inclusion with Product Name, Image, Description, and Specifications
                    inclusion_parts = []
                    
                    # Product Name (this will be the line we search for)
                    if fixture.get('name'):
                        inclusion_parts.append(f"Product Name: {fixture['name']}")
                    
                    # Product Image will be inserted after Product Name (handled separately)
                    
                    # Description
                    if fixture.get('desc'):
                        inclusion_parts.append(f"Description: {fixture['desc']}")
                    
                    # Additional details
                    inclusion_parts.append("(as per annexure)")
                    
                    if fixture.get('hsn'):
                        inclusion_parts.append(f"HSN Code: {fixture['hsn']}")
                    
                    # Specifications
                    if fixture.get('specifications', '').strip():
                        inclusion_parts.append(f"Specifications: {fixture['specifications']}")
                    
                    # Combine all parts for this fixture - EACH FIXTURE GETS ITS OWN BLOCK
                    inclusion_text = f"{idx}. " + "\n   ".join(inclusion_parts)
                    inclusion_items.append({
                        'text': inclusion_text,
                        'image': fixture.get('image') if fixture.get('has_image') or fixture.get('image') else None,
                        'fixture_name': fixture.get('name', f'Fixture {idx}')
                    })
                    
                    # Scope items (simpler format)
                    scope_items.append(f"{idx}. {fixture['name']} - {fixture['desc']}")
                
                # Specifications for separate section
                if fixture.get('specifications', '').strip():
                    spec_item = f"{idx}. {fixture['name']}:\n{fixture['specifications']}"
                    specification_items.append(spec_item)
            
            # Build final content strings - SEPARATE EACH FIXTURE WITH DOUBLE LINE BREAKS
            inclusion_content = "\n\n".join([item['text'] for item in inclusion_items]) if inclusion_items else "Fixtures as per discussion"
            scope_content = "\n".join(scope_items) if scope_items else "Design & manufacturing of fixtures as per discussion"
            specification_content = "\n\n".join(specification_items) if specification_items else "Specifications as per discussion and drawings."
            
            # Prepare fixtures data for database (without image bytes)
            # Note: Image bytes cannot be stored in JSONField, so we only store metadata
            # Images are processed during document generation but not persisted
            fixtures_for_db = []
            for fixture in fixtures:
                db_fixture = dict(fixture)
                # Remove image bytes before saving to database
                if 'image' in db_fixture:
                    del db_fixture['image']
                fixtures_for_db.append(db_fixture)
            
            # Check if there's an existing draft for this quote number and convert it
            existing_draft = None
            if quote_data['quote_no']:
                existing_draft = Quotation.objects.filter(
                    quote_number=quote_data['quote_no'],
                    status='draft',
                    created_by=request.user
                ).first()
            
            # Auto-increment revision for all quotations (always automatic)
            if existing_draft and existing_draft.status == 'draft':
                # Converting draft to final - keep current revision
                new_revision = existing_draft.revision
                print(f"Converting draft to final with revision: {new_revision}")
            else:
                # Check if this is a regeneration of an existing quotation
                # Look for quotations with the same base quote number (without revision suffix)
                base_quote_no = quote_data['quote_no'].split('_')[0]  # Remove any existing suffix
                existing_quotations = Quotation.objects.filter(
                    quote_number__startswith=base_quote_no,
                    created_by=request.user
                ).exclude(status='draft').order_by('-created_at')
                
                if existing_quotations.exists():
                    # This is a regeneration - increment revision and create new quote number
                    latest_quotation = existing_quotations.first()
                    new_revision = increment_revision(latest_quotation.revision)
                    # Create new quote number with revision suffix to avoid UNIQUE constraint
                    new_quote_number = f"{base_quote_no}_{new_revision.replace(' ', '').replace('Rev', 'R')}"
                    print(f"Regenerating: {latest_quotation.quote_number} ({latest_quotation.revision}) → {new_quote_number} ({new_revision})")
                    quote_data['quote_no'] = new_quote_number
                else:
                    # New quotation - always start with Rev A
                    new_revision = 'Rev A'
                    print(f"New quotation with revision: {new_revision}")
            
            # Always use auto-managed revision (ignore user input)
            quote_data['revision'] = new_revision
            
            if existing_draft:
                # Update existing draft to final status instead of creating new
                existing_draft.revision = quote_data['revision']
                existing_draft.quotation_date = quote_data['date']
                existing_draft.to_person = quote_data['to_person']
                existing_draft.firm = quote_data['firm']
                existing_draft.address = quote_data['address']
                existing_draft.payment_terms = quote_data['payment_terms']
                existing_draft.delivery_terms = quote_data['delivery_terms']
                existing_draft.scope_description = scope_content
                existing_draft.fixtures_data = fixtures_for_db
                existing_draft.fixtures_count = len(fixtures)
                existing_draft.status = 'generated'  # Convert draft to final
                existing_draft.save()
                quotation = existing_draft
            else:
                # Save to database as new quotation
                quotation = Quotation.objects.create(
                    quote_number=quote_data['quote_no'],
                    revision=quote_data['revision'],
                    quotation_date=quote_data['date'],
                    to_person=quote_data['to_person'],
                    firm=quote_data['firm'],
                    address=quote_data['address'],
                    payment_terms=quote_data['payment_terms'],
                    delivery_terms=quote_data['delivery_terms'],
                    scope_description=scope_content,
                    fixtures_data=fixtures_for_db,
                    fixtures_count=len(fixtures),
                    status='generated',
                    created_by=request.user
                )
            
            # Load template
            from pathlib import Path
            BASE_DIR = Path(__file__).resolve().parent.parent
            template_path = os.path.join(BASE_DIR, 'Quote KEC005JN2025 RevA - new format full_1.docx')
            
            # Check if template exists, if not create a basic one
            if not os.path.exists(template_path):
                # Create a basic template
                doc = Document()
                
                # Add basic structure
                doc.add_heading('QUOTATION', 0)
                
                # Basic info
                info_p = doc.add_paragraph()
                info_p.add_run('Quote Number: ').bold = True
                info_p.add_run('KEC005JN2025\n')
                info_p.add_run('Revision: ').bold = True
                info_p.add_run('Rev A\n')
                info_p.add_run('Date: ').bold = True
                info_p.add_run('Wednesday, September 24, 2025\n')
                
                # Customer info
                doc.add_heading('Customer Information', level=1)
                customer_p = doc.add_paragraph()
                customer_p.add_run('To: ').bold = True
                customer_p.add_run('Mr. Mohak Dholakia\n')
                customer_p.add_run('Company: ').bold = True
                customer_p.add_run('Schneider Electric India Private Limited\n')
                customer_p.add_run('Address: ').bold = True
                customer_p.add_run('Electrical & Automation (E&A), Village Ankhol, Behind L&T Knowledge City, N. H. 8, Between Ajwa-Waghodia Junction, Vadodara –390019. India\n')
                
                # Terms
                doc.add_heading('Terms & Conditions', level=1)
                terms_p = doc.add_paragraph()
                terms_p.add_run('Payment Terms: ').bold = True
                terms_p.add_run('Payment: 45 Days from delivery (Being an MSME)\n')
                terms_p.add_run('Delivery Terms: ').bold = True
                terms_p.add_run('Delivery: 2-3 weeks per fixture from Purchase Order (Gov. force majeure conditions apply at time of delivery)\n')
                
                # Content sections
                doc.add_heading('Inclusion', level=1)
                doc.add_paragraph('<inclusion>')
                
                doc.add_heading('Scope of Work', level=1)
                doc.add_paragraph('<scope>')
                
                doc.add_heading('Specifications', level=1)
                doc.add_paragraph('<specification>')
                
                # Basic pricing table
                doc.add_heading('Pricing', level=1)
                table = doc.add_table(rows=5, cols=6)
                table.style = 'Table Grid'
                
                # Header
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Sr.'
                hdr_cells[1].text = 'Description'
                hdr_cells[2].text = 'Qty'
                hdr_cells[3].text = 'Rate'
                hdr_cells[4].text = 'Unit'
                hdr_cells[5].text = 'Amount'
                
                # Sample rows
                row1 = table.rows[1].cells
                row1[0].text = '1'
                row1[1].text = 'Fixture 1 Holding minitop connector\n(as per annexure)\nHSN Code: 84790000'
                row1[2].text = '1'
                row1[3].text = '26,879/-'
                row1[4].text = 'Set'
                row1[5].text = '26,879/-'
                
                row2 = table.rows[2].cells
                for cell in row2.cells:
                    cell.text = 'In Words: Twenty-Six Thousand Eight Hundred Seventy-Nine INR Only PER EACH'
                
                row3 = table.rows[3].cells
                row3[0].text = '2'
                row3[1].text = 'Fixture 2 pulling minitop PCBA\n(as per annexure)\nHSN Code: 84790000'
                row3[2].text = '1'
                row3[3].text = '29,546/-'
                row3[4].text = 'Set'
                row3[5].text = '29,546/-'
                
                row4 = table.rows[4].cells
                for cell in row4.cells:
                    cell.text = 'In Words: Twenty-Nine Thousand Five Hundred Forty-Six INR Only PER EACH'
                
                print(f"Created basic template at: {template_path}")
            else:
                doc = Document(template_path)
            
            # Prepare replacements for tag-based system
            replacements = {
                # Basic document info
                "KEC005JN2025": quote_data['quote_no'],
                "Rev A": quote_data['revision'],
                "Wednesday, September 24, 2025": quote_data['date'],
                "Mr. Mohak Dholakia": quote_data['to_person'],
                "Schneider Electric India Private Limited": quote_data['firm'],
                "Electrical & Automation (E&A), Village Ankhol, Behind L&T Knowledge City, N. H. 8, Between Ajwa-Waghodia Junction, Vadodara –390019. India": quote_data['address'],
                
                # Terms
                "Payment: 45 Days from delivery (Being an MSME)": quote_data['payment_terms'],
                "Delivery: 2-3 weeks per fixture from Purchase Order (Gov. force majeure conditions apply at time of delivery)": quote_data['delivery_terms'],
                
                # Tag-based content sections
                "<inclusion>": inclusion_content,
                "<scope>": scope_content,
                "<specification>": specification_content,
                "< specification>": specification_content,  # Handle the space variant
                
                # Price replacements only (avoid fixture name replacements that interfere with Sr column)
                "26,879/-": fixtures[0]['price'] if fixtures[0]['price'] else "26,879/-",
                "Twenty-Six Thousand Eight Hundred Seventy-Nine INR Only PER EACH": fixtures[0]['words'] if fixtures[0]['words'] else "Twenty-Six Thousand Eight Hundred Seventy-Nine INR Only PER EACH",
                
                "29,546/-": fixtures[1]['price'] if fixtures[1]['price'] else "29,546/-",
                "Twenty-Nine Thousand Five Hundred Forty-Six INR Only PER EACH": fixtures[1]['words'] if fixtures[1]['words'] else "Twenty-Nine Thousand Five Hundred Forty-Six INR Only PER EACH",
            }
            
            # Apply text replacements using the advanced seek_and_replace function
            doc = seek_and_replace(doc, replacements)
            
            # Insert images in inclusion section after text replacements
            insert_images_in_inclusion_section(doc, fixtures)
            
            # CRITICAL: Handle pricing table AFTER all text replacements to prevent interference
            print("=== STARTING TABLE PROCESSING ===")
            if len(doc.tables) > 1:
                pricing_table = doc.tables[1]  # Second table is the pricing table
                print(f"Found pricing table with {len(pricing_table.rows)} rows")
                
                # AGGRESSIVE: Populate existing fixture rows with forced content
                populate_pricing_table_with_fixtures(pricing_table, fixtures)
                
                # Add additional fixture rows if needed (beyond the first 2)
                if len(fixtures) > 2:
                    print(f"Adding {len(fixtures) - 2} additional fixture rows")
                    add_fixture_rows_to_table(pricing_table, fixtures)
                    
                # FINAL VERIFICATION: Check if our changes took effect
                print("=== POST-PROCESSING VERIFICATION ===")
                for r_idx in range(min(6, len(pricing_table.rows))):
                    try:
                        row = pricing_table.rows[r_idx]
                        first_cell = row.cells[0].text.strip() if len(row.cells) > 0 else "N/A"
                        second_cell = row.cells[1].text.strip()[:50] + "..." if len(row.cells) > 1 and len(row.cells[1].text.strip()) > 50 else row.cells[1].text.strip() if len(row.cells) > 1 else "N/A"
                        print(f"Final Row {r_idx}: Sr='{first_cell}' | Desc='{second_cell}'")
                    except Exception as e:
                        print(f"Final Row {r_idx}: Error - {e}")
                        
            else:
                print("WARNING: Could not find pricing table")
            print("=== TABLE PROCESSING COMPLETE ===")
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Update quotation status
            quotation.status = 'generated'
            quotation.save()
            
            # Store the file data in session for download
            request.session['download_file'] = {
                'content': base64.b64encode(buffer.getvalue()).decode('utf-8'),
                'filename': f"Quotation_{quote_data['quote_no']}.docx",
                'quotation_id': quotation.id
            }
            
            messages.success(request, f'Quotation {quotation.quote_number} generated successfully!')
            
            # Return a special response that triggers download then redirect
            return render(request, 'dashboard/quotation_download_redirect.html', {
                'quotation_id': quotation.id,
                'filename': f"Quotation_{quote_data['quote_no']}.docx"
            })
            
        except Exception as e:
            messages.error(request, f'Error generating quotation: {str(e)}')
            import traceback
            traceback.print_exc()
            return render(request, 'dashboard/quotation_generator.html', {'error': str(e)})
    
    return render(request, 'dashboard/quotation_generator.html')

@login_required
def save_quotation_draft(request):
    """Save quotation data as draft without generating document"""
    if request.method == 'POST':
        try:
            print("=== DRAFT SAVE DEBUG ===")
            print(f"POST data keys: {list(request.POST.keys())}")
            print(f"FILES data keys: {list(request.FILES.keys())}")
            
            # Collect form data (minimal validation for draft)
            quote_data = {
                'quote_no': request.POST.get('quote_no', ''),
                'revision': request.POST.get('revision', 'Rev A'),
                'date': request.POST.get('date', ''),
                'to_person': request.POST.get('to_person', ''),
                'firm': request.POST.get('firm', ''),
                'address': request.POST.get('address', ''),
                'payment_terms': request.POST.get('payment_terms', ''),
                'delivery_terms': request.POST.get('delivery_terms', ''),
            }
            
            # Process fixtures data (allow incomplete data for draft)
            fixtures = []
            fixture_index = 0
            
            while True:
                name_key = f'fixtures[{fixture_index}][name]'
                if name_key not in request.POST:
                    break
                    
                fixture = {
                    'name': request.POST.get(name_key, ''),
                    'desc': request.POST.get(f'fixtures[{fixture_index}][desc]', ''),
                    'hsn': request.POST.get(f'fixtures[{fixture_index}][hsn]', ''),
                    'qty': request.POST.get(f'fixtures[{fixture_index}][qty]', ''),
                    'unit': request.POST.get(f'fixtures[{fixture_index}][unit]', ''),
                    'price': request.POST.get(f'fixtures[{fixture_index}][price]', ''),
                    'total': request.POST.get(f'fixtures[{fixture_index}][total]', ''),
                    'words': request.POST.get(f'fixtures[{fixture_index}][words]', ''),
                    'specifications': request.POST.get(f'fixtures[{fixture_index}][specifications]', ''),
                }
                fixtures.append(fixture)
                fixture_index += 1
            
            # Check if draft already exists for this user and quote number
            existing_draft = None
            if quote_data['quote_no']:
                existing_draft = Quotation.objects.filter(
                    quote_number=quote_data['quote_no'],
                    status='draft',
                    created_by=request.user
                ).first()
            
            if existing_draft:
                # Update existing draft
                existing_draft.revision = quote_data['revision']
                existing_draft.quotation_date = quote_data['date']
                existing_draft.to_person = quote_data['to_person']
                existing_draft.firm = quote_data['firm']
                existing_draft.address = quote_data['address']
                existing_draft.payment_terms = quote_data['payment_terms']
                existing_draft.delivery_terms = quote_data['delivery_terms']
                existing_draft.fixtures_data = fixtures
                existing_draft.fixtures_count = len(fixtures)
                existing_draft.save()
                
                # Handle image uploads for existing draft
                from .models import DraftImage
                
                # Clear existing images for this draft
                existing_draft.draft_images.all().delete()
                
                # Save images with error handling
                for idx, fixture in enumerate(fixtures):
                    image_key = f'fixtures[{idx}][image]'
                    print(f"Processing image for fixture {idx}, key: {image_key}")
                    if image_key in request.FILES:
                        uploaded_file = request.FILES[image_key]
                        print(f"Found uploaded file: {uploaded_file.name}, size: {uploaded_file.size}")
                        if uploaded_file and uploaded_file.size > 0:
                            try:
                                # Create DraftImage record
                                draft_image = DraftImage.objects.create(
                                    quotation=existing_draft,
                                    fixture_index=idx,
                                    image=uploaded_file,
                                    original_filename=uploaded_file.name,
                                    file_size=uploaded_file.size
                                )
                                # Update fixture data to indicate image exists
                                fixtures[idx]['has_image'] = True
                                print(f"✅ Saved draft image: {uploaded_file.name} for fixture {idx}, ID: {draft_image.id}")
                            except Exception as img_error:
                                print(f"❌ Error saving image for fixture {idx}: {img_error}")
                                import traceback
                                traceback.print_exc()
                                # Don't fail the entire draft save for image errors
                                fixtures[idx]['has_image'] = False
                    else:
                        print(f"No image file found for fixture {idx}")
                
                # Update fixtures data with image info
                existing_draft.fixtures_data = fixtures
                existing_draft.save()
                
                from django.contrib import messages
                from django.shortcuts import redirect
                messages.success(request, 'Draft updated successfully!')
                return redirect('dashboard:quotation_management')
            else:
                # Create new draft
                from django.utils import timezone
                draft = Quotation.objects.create(
                    quote_number=quote_data['quote_no'] or f"DRAFT_{request.user.id}_{timezone.now().strftime('%Y%m%d_%H%M%S')}",
                    revision=quote_data['revision'],
                    quotation_date=quote_data['date'],
                    to_person=quote_data['to_person'],
                    firm=quote_data['firm'],
                    address=quote_data['address'],
                    payment_terms=quote_data['payment_terms'],
                    delivery_terms=quote_data['delivery_terms'],
                    fixtures_data=fixtures,
                    fixtures_count=len(fixtures),
                    status='draft',  # EXPLICIT DRAFT STATUS
                    created_by=request.user
                )
                
                # Handle image uploads for new draft
                from .models import DraftImage
                
                # Save images with error handling
                for idx, fixture in enumerate(fixtures):
                    image_key = f'fixtures[{idx}][image]'
                    print(f"Processing image for fixture {idx}, key: {image_key}")
                    if image_key in request.FILES:
                        uploaded_file = request.FILES[image_key]
                        print(f"Found uploaded file: {uploaded_file.name}, size: {uploaded_file.size}")
                        if uploaded_file and uploaded_file.size > 0:
                            try:
                                # Create DraftImage record
                                draft_image = DraftImage.objects.create(
                                    quotation=draft,
                                    fixture_index=idx,
                                    image=uploaded_file,
                                    original_filename=uploaded_file.name,
                                    file_size=uploaded_file.size
                                )
                                # Update fixture data to indicate image exists
                                fixtures[idx]['has_image'] = True
                                print(f"✅ Saved draft image: {uploaded_file.name} for fixture {idx}, ID: {draft_image.id}")
                            except Exception as img_error:
                                print(f"❌ Error saving image for fixture {idx}: {img_error}")
                                import traceback
                                traceback.print_exc()
                                # Don't fail the entire draft save for image errors
                                fixtures[idx]['has_image'] = False
                    else:
                        print(f"No image file found for fixture {idx}")
                
                # Update fixtures data with image info
                draft.fixtures_data = fixtures
                draft.save()
                
                from django.contrib import messages
                from django.shortcuts import redirect
                messages.success(request, 'Draft saved successfully!')
                return redirect('dashboard:quotation_management')
                
        except Exception as e:
            from django.contrib import messages
            from django.shortcuts import redirect
            messages.error(request, f'Error saving draft: {str(e)}')
            return redirect('dashboard:quotation_generator')
    
    from django.shortcuts import redirect
    return redirect('dashboard:quotation_generator')

@login_required
def quotation_management_view(request):
    """Quotation management page with list of quotations grouped by base quote number, showing latest revision"""
    search_query = request.GET.get('search', '')
    
    # Get all quotations and group by base quote number (without revision suffix)
    from django.db.models import Max
    
    # First get all quotations
    all_quotations = Quotation.objects.all()
    
    # Apply search filter if provided
    if search_query:
        all_quotations = all_quotations.filter(
            models.Q(quote_number__icontains=search_query) |
            models.Q(firm__icontains=search_query) |
            models.Q(to_person__icontains=search_query)
        )
    
    # Group quotations by base quote number and get latest revision for each
    quotation_groups = {}
    
    for quotation in all_quotations:
        # Extract base quote number (remove revision suffix like _RA, _RB)
        base_quote = quotation.quote_number.split('_')[0]
        
        # Keep only the latest updated quotation for each base quote number
        if base_quote not in quotation_groups:
            quotation_groups[base_quote] = quotation
        else:
            # Compare updated_at timestamps and keep the latest
            if quotation.updated_at > quotation_groups[base_quote].updated_at:
                quotation_groups[base_quote] = quotation
    
    # Convert to list and sort by updated_at (most recent first)
    quotations_list = list(quotation_groups.values())
    quotations_list.sort(key=lambda x: x.updated_at, reverse=True)
    
    # Pagination
    paginator = Paginator(quotations_list, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Calculate statistics - based on unique base quote numbers
    total_count = len(quotation_groups)  # Unique quotations (latest revisions only)
    generated_count = sum(1 for q in quotation_groups.values() if q.status == 'generated')
    draft_count = sum(1 for q in quotation_groups.values() if q.status == 'draft')
    this_month_count = sum(1 for q in quotation_groups.values() 
                          if q.updated_at.month == datetime.now().month and 
                             q.updated_at.year == datetime.now().year)
    
    context = {
        'quotations': page_obj,
        'search_query': search_query,
        'total_count': total_count,
        'generated_count': generated_count,
        'draft_count': draft_count,
        'this_month_count': this_month_count,
    }
    
    return render(request, 'dashboard/quotation_management.html', context)

@login_required
def quotation_edit(request, quotation_id):
    """Edit existing quotation"""
    quotation = get_object_or_404(Quotation, id=quotation_id)
    
    if request.method == 'POST':
        try:
            # Update quotation data
            quotation.quote_number = request.POST.get('quote_no', quotation.quote_number)
            quotation.revision = request.POST.get('revision', quotation.revision)
            quotation.quotation_date = request.POST.get('date', quotation.quotation_date)
            quotation.to_person = request.POST.get('to_person', quotation.to_person)
            quotation.firm = request.POST.get('firm', quotation.firm)
            quotation.address = request.POST.get('address', quotation.address)
            quotation.payment_terms = request.POST.get('payment_terms', quotation.payment_terms)
            quotation.delivery_terms = request.POST.get('delivery_terms', quotation.delivery_terms)
            quotation.scope_description = request.POST.get('scope_desc', quotation.scope_description)
            quotation.scope_line_1 = request.POST.get('scope_1', quotation.scope_line_1)
            quotation.scope_line_2 = request.POST.get('scope_2', quotation.scope_line_2)
            
            # Process fixtures
            fixtures_data = []
            fixture_index = 0
            
            while True:
                name_key = f'fixtures[{fixture_index}][name]'
                if name_key not in request.POST:
                    break
                    
                fixture = {
                    'name': request.POST.get(name_key, ''),
                    'desc': request.POST.get(f'fixtures[{fixture_index}][desc]', ''),
                    'hsn': request.POST.get(f'fixtures[{fixture_index}][hsn]', ''),
                    'qty': request.POST.get(f'fixtures[{fixture_index}][qty]', ''),
                    'unit': request.POST.get(f'fixtures[{fixture_index}][unit]', ''),
                    'price': request.POST.get(f'fixtures[{fixture_index}][price]', ''),
                    'total': request.POST.get(f'fixtures[{fixture_index}][total]', ''),
                    'words': request.POST.get(f'fixtures[{fixture_index}][words]', ''),
                    'specifications': request.POST.get(f'fixtures[{fixture_index}][specifications]', ''),
                    'has_image': False,  # For edit, we don't handle image uploads currently
                }
                fixtures_data.append(fixture)
                fixture_index += 1
            
            quotation.fixtures_data = fixtures_data
            quotation.fixtures_count = len(fixtures_data)
            quotation.save()
            
            messages.success(request, f'Quotation {quotation.quote_number} updated successfully!')
            return redirect('dashboard:quotation_management')
            
        except Exception as e:
            messages.error(request, f'Error updating quotation: {str(e)}')
    
    context = {
        'quotation': quotation,
        'is_edit': True,
        'title': 'Edit Quotation',
        'action': 'Edit',
    }
    return render(request, 'dashboard/quotation_form.html', context)

@login_required
@login_required
def quotation_download_file(request):
    """Serve the generated quotation file from session"""
    if 'download_file' in request.session:
        file_data = request.session['download_file']
        
        # Decode the base64 content
        file_content = base64.b64decode(file_data['content'])
        
        # Clear the session data
        del request.session['download_file']
        
        # Return the file
        response = HttpResponse(
            file_content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{file_data["filename"]}"'
        return response
    
    # If no file in session, redirect to management
    return redirect('dashboard:quotation_management')

@login_required
def quotation_download(request, quotation_id):
    """Download latest revision of quotation with proper functionality"""
    quotation = get_object_or_404(Quotation, id=quotation_id)
    
    # Get the base quote number and find the latest revision
    base_quote_no = quotation.quote_number.split('_')[0]
    
    # Find the latest updated quotation with the same base quote number
    latest_quotation = Quotation.objects.filter(
        quote_number__startswith=base_quote_no,
        created_by=quotation.created_by
    ).order_by('-updated_at').first()
    
    # Use the latest quotation for download
    if latest_quotation:
        quotation = latest_quotation
        print(f"Downloading latest revision: {quotation.quote_number} (updated: {quotation.updated_at})")
    
    try:
        # Always regenerate the document (no file storage)
        # Prepare replacements for document generation
        replacements = {
            # Basic Info
            "KEC005JN2025": quotation.quote_number,
            "Rev A": quotation.revision,
            "Wednesday, September 24, 2025": quotation.quotation_date,
            "Mr. Mohak Dholakia": quotation.to_person,
            "Schneider Electric India Private Limited": quotation.firm,
            "Electrical & Automation (E&A), Village Ankhol, Behind L&T Knowledge City, N. H. 8, Between Ajwa-Waghodia Junction, Vadodara –390019. India": quotation.address,
            
            # Payment and Delivery
            "Payment: 45 Days from delivery (Being an MSME)": quotation.payment_terms,
            "Delivery: 2-3 weeks per fixture from Purchase Order (Gov. force majeure conditions apply at time of delivery)": quotation.delivery_terms,
            
            # Scope
            "Fixtures as per discussion:": quotation.scope_description,
        }

        # Build specifications from stored fixtures data
        specification_items = []
        for idx, fixture in enumerate(quotation.fixtures_data, 1):
            if fixture.get('specifications', '').strip():
                spec_item = f"{idx}. {fixture.get('name', '')}:\n{fixture.get('specifications', '')}"
                specification_items.append(spec_item)
        specifications_str = "\n\n".join(specification_items) if specification_items else "Specifications as per discussion and drawings."
        
        # Build enhanced inclusion content from stored fixtures data
        inclusion_items = []
        for idx, fixture in enumerate(quotation.fixtures_data, 1):
            if fixture.get('name') or fixture.get('desc'):
                # Enhanced inclusion with Product Name, Description, and Specifications
                inclusion_parts = []
                
                # Product Name
                if fixture.get('name'):
                    inclusion_parts.append(f"Product Name: {fixture['name']}")
                
                # Product Image will be inserted after Product Name (handled separately)
                
                # Description
                if fixture.get('desc'):
                    inclusion_parts.append(f"Description: {fixture['desc']}")
                
                # Additional details
                inclusion_parts.append("(as per annexure)")
                
                if fixture.get('hsn'):
                    inclusion_parts.append(f"HSN Code: {fixture['hsn']}")
                
                # Specifications
                if fixture.get('specifications', '').strip():
                    inclusion_parts.append(f"Specifications: {fixture['specifications']}")
                
                # Note: Images are not available for download (no image bytes stored)
                if fixture.get('has_image'):
                    inclusion_parts.append("Product Image: (Image was included in original quotation)")
                
                # Combine all parts for this fixture
                inclusion_text = f"{idx}. " + "\n   ".join(inclusion_parts)
                inclusion_items.append(inclusion_text)
        
        inclusion_content = "\n\n".join(inclusion_items) if inclusion_items else "Fixtures as per discussion"
        
        # Add inclusion content to replacements
        replacements["<inclusion>"] = inclusion_content
        
        # Add tag-based replacements for specifications
        replacements["<specification>"] = specifications_str
        replacements["< specification>"] = specifications_str  # Handle space variant

        # Add fixture replacements for first 2 fixtures (prices and words only, not names)
        if len(quotation.fixtures_data) >= 1:
            fixture = quotation.fixtures_data[0]
            replacements["26,879/-"] = fixture.get('price', '26,879/-')
            replacements["Twenty-Six Thousand Eight Hundred Seventy-Nine INR Only PER EACH"] = fixture.get('words', 'Twenty-Six Thousand Eight Hundred Seventy-Nine INR Only PER EACH')
        
        if len(quotation.fixtures_data) >= 2:
            fixture = quotation.fixtures_data[1]
            replacements["29,546/-"] = fixture.get('price', '29,546/-')
            replacements["Twenty-Nine Thousand Five Hundred Forty-Six INR Only PER EACH"] = fixture.get('words', 'Twenty-Nine Thousand Five Hundred Forty-Six INR Only PER EACH')

        # Load template
        from pathlib import Path
        BASE_DIR = Path(__file__).resolve().parent.parent
        template_path = os.path.join(BASE_DIR, 'Quote KEC005JN2025 RevA - new format full_1.docx')
        
        doc = Document(template_path)

        # Apply text replacements using the advanced seek_and_replace function
        doc = seek_and_replace(doc, replacements)

        # Handle pricing table with fixtures (without images for download)
        if len(doc.tables) > 1:
            pricing_table = doc.tables[1]  # Second table is the pricing table
            
            # For download, create fixtures without image bytes
            fixtures_for_doc = []
            for fixture in quotation.fixtures_data:
                doc_fixture = dict(fixture)
                doc_fixture['image'] = None  # No image bytes available for download
                doc_fixture['has_image'] = False
                fixtures_for_doc.append(doc_fixture)
            
            # Populate existing fixture rows (first 2)
            populate_pricing_table_with_fixtures(pricing_table, fixtures_for_doc)
            
            # Add additional fixture rows if needed
            if len(fixtures_for_doc) > 2:
                add_fixture_rows_to_table(pricing_table, fixtures_for_doc)

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Return as download
        response = HttpResponse(
            buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="Quotation_{quotation.quote_number}.docx"'
        return response

    except Exception as e:
        messages.error(request, f'Error downloading quotation: {str(e)}')
        import traceback
        traceback.print_exc()
        return redirect('dashboard:quotation_management')

@login_required
def process_po_pdf_ajax(request):
    """
    AJAX endpoint to process uploaded PDF and extract form data using Mistral AI
    """
    if request.method == 'POST' and request.FILES.get('pdf_file'):
        try:
            from .services import extract_po_data_from_pdf
            
            pdf_file = request.FILES['pdf_file']
            
            # Validate file type
            if not pdf_file.name.lower().endswith('.pdf'):
                return JsonResponse({
                    'success': False,
                    'error': 'Please upload a PDF file only.'
                })
            
            # Validate file size (5MB limit)
            if pdf_file.size > 5 * 1024 * 1024:
                return JsonResponse({
                    'success': False,
                    'error': 'File size must be less than 5MB.'
                })
            
            # Extract data using Mistral AI
            extracted_data = extract_po_data_from_pdf(pdf_file)
            
            # Check if extraction failed
            if extracted_data.get('error'):
                return JsonResponse({
                    'success': False,
                    'error': extracted_data['error']
                })
            
            # Map extracted data to form fields
            form_data = {
                'po_number': extracted_data.get('po_number', ''),
                'order_date': extracted_data.get('po_date', ''),  # po_date from OCR maps to order_date in form
                'customer_name': extracted_data.get('customer_name', ''),
                'order_value': extracted_data.get('net_value', ''),  # net_value from OCR maps to order_value in form
                'remarks': extracted_data.get('remarks', ''),
            }
            
            return JsonResponse({
                'success': True,
                'data': form_data,
                'message': 'PDF processed successfully!'
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': f'Error processing PDF: {str(e)}'
            })
    
    return JsonResponse({
        'success': False,
        'error': 'No PDF file uploaded.'
    })

# Demo and API Views for Searchable Input Component

@login_required
def searchable_input_demo_view(request):
    """Demo page for searchable input component"""
    return render(request, 'dashboard/searchable_input_demo.html', {
        'title': 'Searchable Input Demo'
    })

def get_user_names_api(request):
    """API endpoint to get user names for searchable input"""
    try:
        # Get all user names (first name + last name or username)
        users = User.objects.all()
        user_names = []
        
        for user in users:
            full_name = user.get_full_name()
            if full_name.strip():
                user_names.append(full_name)
            else:
                user_names.append(user.username)
        
        return JsonResponse(user_names, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def get_contact_names_api(request):
    """API endpoint to get contact names for searchable input"""
    try:
        # Get all contact names from the Contact model
        contacts = Contact.objects.all().values_list('contact_name', flat=True)
        contact_names = list(contacts)
        
        return JsonResponse(contact_names, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@login_required
def searchable_integration_demo_view(request):
    """Integration demo page for searchable input component"""
    return render(request, 'dashboard/searchable_integration_example.html', {
        'title': 'Searchable Input Integration Demo'
    })
def get_company_names_api(request):
    """API endpoint to get company names for searchable input"""
    try:
        # Get all company names from the Company model
        companies = Company.objects.all()
        company_names = []
        
        for company in companies:
            company_names.append(f"{company.company_name} - {company.city}")
        
        return JsonResponse(company_names, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def get_inquiry_search_data_api(request):
    """API endpoint to get inquiry search data for searchable input"""
    try:
        # Get inquiry data for search
        inquiries = InquiryHandler.objects.all()
        search_data = []
        
        for inquiry in inquiries:
            # Add quote number
            if inquiry.quote_no:
                search_data.append(inquiry.quote_no)
            
            # Add company name
            if inquiry.company and inquiry.company.contact_name:
                search_data.append(inquiry.company.contact_name)
            
            # Add customer name
            if inquiry.customer_name:
                search_data.append(inquiry.customer_name)
            
            # Add sales person name
            if inquiry.sales:
                search_data.append(inquiry.sales.get_full_name() or inquiry.sales.username)
        
        # Remove duplicates and return
        unique_data = list(set(search_data))
        return JsonResponse(unique_data, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def get_quotation_search_data_api(request):
    """API endpoint to get quotation search data for searchable input"""
    try:
        # Get quotation data for search
        quotations = Quotation.objects.all()
        search_data = []
        
        for quotation in quotations:
            # Add quote number
            if quotation.quote_number:
                search_data.append(quotation.quote_number)
            
            # Add firm name
            if quotation.firm:
                search_data.append(quotation.firm)
            
            # Add to person
            if quotation.to_person:
                search_data.append(quotation.to_person)
        
        # Remove duplicates and return
        unique_data = list(set(search_data))
        return JsonResponse(unique_data, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def get_purchase_order_search_data_api(request):
    """API endpoint to get purchase order search data for searchable input"""
    try:
        # Get purchase order data for search
        orders = PurchaseOrder.objects.all()
        search_data = []
        
        for order in orders:
            # Add PO number
            if order.po_number:
                search_data.append(order.po_number)
            
            # Add customer name
            if order.customer_name:
                search_data.append(order.customer_name)
            
            # Add company name
            if order.company and order.company.contact_name:
                search_data.append(order.company.contact_name)
            
            # Add sales person
            if order.sales_person:
                search_data.append(order.sales_person.get_full_name() or order.sales_person.username)
            
            # Add project manager
            if order.project_manager:
                search_data.append(order.project_manager.get_full_name() or order.project_manager.username)
        
        # Remove duplicates and return
        unique_data = list(set(search_data))
        return JsonResponse(unique_data, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def get_invoice_search_data_api(request):
    """API endpoint to get invoice search data for searchable input"""
    try:
        # Get invoice data for search
        invoices = Invoice.objects.all()
        search_data = []
        
        for invoice in invoices:
            # Add invoice number
            if invoice.invoice_number:
                search_data.append(invoice.invoice_number)
            
            # Add customer name
            if invoice.customer_name:
                search_data.append(invoice.customer_name)
            
            # Add company name
            if invoice.company and invoice.company.contact_name:
                search_data.append(invoice.company.contact_name)
        
        # Remove duplicates and return
        unique_data = list(set(search_data))
        return JsonResponse(unique_data, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
def get_additional_supply_search_data_api(request):
    """API endpoint to get additional supply search data for searchable input"""
    try:
        # Get additional supply data for search
        supplies = AdditionalSupply.objects.all()
        search_data = []
        
        for supply in supplies:
            # Add invoice number
            if supply.invoice and supply.invoice.invoice_number:
                search_data.append(supply.invoice.invoice_number)
            
            # Add customer name
            if supply.invoice and supply.invoice.customer_name:
                search_data.append(supply.invoice.customer_name)
            
            # Add company name
            if supply.invoice and supply.invoice.company and supply.invoice.company.contact_name:
                search_data.append(supply.invoice.company.contact_name)
            
            # Add description
            if supply.description:
                search_data.append(supply.description)
        
        # Remove duplicates and return
        unique_data = list(set(search_data))
        return JsonResponse(unique_data, safe=False)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
@login_required
def search_fix_demo_view(request):
    """Demo page showing the search filter UI fix"""
    return render(request, 'dashboard/search_fix_demo.html', {
        'title': 'Search Filter UI Fix Demo'
    })
@login_required
def search_test_demo_view(request):
    """Test page for the fixed search dropdown"""
    return render(request, 'dashboard/search_test.html', {
        'title': 'Search Dropdown Test - Fixed Version'
    })
@login_required
def simple_search_test_view(request):
    """Simple search test page with working dropdown fix"""
    return render(request, 'dashboard/simple_search_test.html', {
        'title': 'Simple Search Test - Working Fix'
    })

@login_required
def notifications_view(request):
    """Display notifications for admin users"""
    from django.utils import timezone
    from datetime import timedelta
    
    # Only show notifications to admin/manager users
    if not (hasattr(request.user, 'userprofile') and request.user.userprofile.get_roles_list() in ['admin', 'manager']):
        messages.error(request, 'Access denied. Only administrators can view notifications.')
        return redirect('dashboard:dashboard')
    
    # Get notifications from the last 48 hours only, ordered by newest first
    cutoff_time = timezone.now() - timedelta(hours=48)
    notifications = Notification.objects.filter(created_at__gte=cutoff_time).order_by('-created_at')
    
    # Mark notifications as read when viewed
    unread_notifications = notifications.filter(is_read=False)
    unread_notifications.update(is_read=True)
    
    return render(request, 'dashboard/notifications.html', {
        'notifications': notifications,
        'total_count': notifications.count(),
    })

@login_required
def get_notification_count(request):
    """AJAX endpoint to get unread notification count"""
    from django.utils import timezone
    from datetime import timedelta
    
    if not (hasattr(request.user, 'userprofile') and request.user.userprofile.get_roles_list() in ['admin', 'manager']):
        return JsonResponse({'count': 0})
    
    # Only count notifications from the last 48 hours
    cutoff_time = timezone.now() - timedelta(hours=48)
    unread_count = Notification.objects.filter(
        is_read=False,
        created_at__gte=cutoff_time
    ).count()
    return JsonResponse({'count': unread_count})

@login_required
def cleanup_old_notifications_ajax(request):
    """AJAX endpoint to clean up notifications older than 48 hours"""
    from django.utils import timezone
    from datetime import timedelta
    
    # Only allow admin/manager users to trigger cleanup
    if not (hasattr(request.user, 'userprofile') and request.user.userprofile.get_roles_list() in ['admin', 'manager']):
        return JsonResponse({'success': False, 'error': 'Access denied'})
    
    try:
        # Calculate the cutoff time (1 minute ago for testing)
        cutoff_time = timezone.now() - timedelta(minutes=1)
        
        # Find and delete old notifications
        old_notifications = Notification.objects.filter(created_at__lt=cutoff_time)
        deleted_count, _ = old_notifications.delete()
        
        return JsonResponse({
            'success': True,
            'deleted_count': deleted_count,
            'message': f'Cleaned up {deleted_count} old notifications'
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })