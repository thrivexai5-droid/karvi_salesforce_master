from django.shortcuts import render
from django.http import HttpResponse
from django.conf import settings
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import os

def seek_and_replace(doc, replacements):
    """Replace text in DOCX while preserving formatting, especially for highlighted text"""
    def replace_in_paragraph(paragraph, replacements):
        # Get full text from all runs
        full_text = "".join(run.text for run in paragraph.runs)
        replaced_text = full_text
        
        # Try replacements on full text
        for old, new in replacements.items():
            if new and old in replaced_text:
                replaced_text = replaced_text.replace(old, new)
        
        # If text changed, rebuild the paragraph
        if replaced_text != full_text:
            # Check if any run has highlighting
            has_highlighted = any(run.font.highlight_color == WD_COLOR_INDEX.YELLOW for run in paragraph.runs)
            
            if has_highlighted:
                # Keep first run's formatting and clear others
                first_run = paragraph.runs[0] if paragraph.runs else None
                if first_run:
                    # Clear all runs
                    for run in paragraph.runs[1:]:
                        run.text = ""
                    # Set new text in first run
                    first_run.text = replaced_text
            else:
                # Non-highlighted text - simple replacement
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = replaced_text

    def replace_in_table(table, replacements):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)
                for inner_table in cell.tables:
                    replace_in_table(inner_table, replacements)

    # Replace in paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)

    # Replace in tables
    for table in doc.tables:
        replace_in_table(table, replacements)

    # Replace in headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_paragraph(para, replacements)
        for para in section.footer.paragraphs:
            replace_in_paragraph(para, replacements)

    return doc

def add_fixture_rows_to_table(table, fixtures_data):
    """Add additional fixture rows to the pricing table"""
    # The table structure: Row 0 is header, Row 1 is Fixture 1, Row 2 is words, Row 3 is Fixture 2, Row 4 is words
    # We need to add fixtures starting from index 2 (3rd fixture)
    if len(fixtures_data) <= 2:
        return  # No additional rows needed

    # Get the template row (use row 3 as template for additional fixtures)
    template_row_index = 3
    
    for idx in range(2, len(fixtures_data)):
        fixture = fixtures_data[idx]
        
        # Add new row for fixture
        new_row = table.add_row()
        cells = new_row.cells
        
        # Fill in the fixture data
        cells[0].text = str(idx + 1)  # Sr number
        cells[1].text = f"{fixture['name']} {fixture['desc']}\n(as per annexure)\nHSN Code: {fixture['hsn']}"
        cells[2].text = fixture['qty']
        cells[3].text = fixture['price']
        cells[4].text = fixture['unit']
        cells[5].text = fixture['total']
        
        # Add words row
        words_row = table.add_row()
        for cell in words_row.cells:
            cell.text = f"In Words: {fixture['words']}"

def quotation_view(request):
    return render(request, 'quotation/index.html')

def generate_quotation(request):
    if request.method == 'POST':
        try:
            # Get form data
            replacements = {
                # Basic Info
                "KEC005JN2025": request.POST.get('quote_no', ''),
                "Rev A": request.POST.get('revision', ''),
                "Wednesday, September 24, 2025": request.POST.get('date', ''),
                "Mr. Mohak Dholakia": request.POST.get('to_person', ''),
                "Schneider Electric India Private Limited": request.POST.get('firm', ''),
                "Electrical & Automation (E&A), Village Ankhol, Behind L&T Knowledge City, N. H. 8, Between Ajwa-Waghodia Junction, Vadodara –390019. India": request.POST.get('address', ''),
                
                # Payment and Delivery
                "Payment: 45 Days from delivery (Being an MSME)": request.POST.get('payment_terms', ''),
                "Delivery: 2-3 weeks per fixture from Purchase Order (Gov. force majeure conditions apply at time of delivery)": request.POST.get('delivery_terms', ''),
                
                # Scope
                "Fixtures as per discussion:": request.POST.get('scope_desc', ''),
                "Design & manufacturing of small size fixtures": request.POST.get('scope_1', ''),
                "Supply": request.POST.get('scope_2', ''),
            }

            # Process dynamic fixtures
            fixtures_data = []
            fixture_index = 0
            
            while True:
                name_key = f'fixtures[{fixture_index}][name]'
                if name_key not in request.POST:
                    break
                    
                fixture = {
                    'name': request.POST.get(name_key, ''),
                    'desc': request.POST.get(f'fixtures[{fixture_index}][desc]', ''),
                    'hsn': request.POST.get(f'fixtures[{fixture_index}][hsn]', ''),
                    'qty': request.POST.get(f'fixtures[{fixture_index}][qty]', ''),
                    'unit': request.POST.get(f'fixtures[{fixture_index}][unit]', ''),
                    'price': request.POST.get(f'fixtures[{fixture_index}][price]', ''),
                    'total': request.POST.get(f'fixtures[{fixture_index}][total]', ''),
                    'words': request.POST.get(f'fixtures[{fixture_index}][words]', ''),
                }
                fixtures_data.append(fixture)
                fixture_index += 1

            # Add fixture replacements based on template defaults
            default_fixtures = [
                {
                    'name': 'Fixture 1 ',
                    'desc': 'Holding minitop connector',
                    'hsn': '84790000',
                    'price': '26,879/-',
                    'words': 'Twenty-Six Thousand Eight Hundred Seventy-Nine INR Only PER EACH'
                },
                {
                    'name': 'Fixture 2 ',
                    'desc': 'pulling minitop PCBA',
                    'hsn': '84790000',
                    'price': '29,546/-',
                    'words': 'Twenty-Nine Thousand Five Hundred Forty-Six INR Only PER EACH'
                }
            ]

            # Map fixtures to replacements
            for idx, fixture in enumerate(fixtures_data):
                if idx < len(default_fixtures):
                    default = default_fixtures[idx]
                    replacements[default['name']] = fixture['name'] + " "
                    replacements[default['desc']] = fixture['desc']
                    replacements[default['price']] = fixture['price']
                    replacements[default['words']] = fixture['words']

            # Load template
            from pathlib import Path
            BASE_DIR = Path(__file__).resolve().parent.parent
            template_path = os.path.join(BASE_DIR, 'Quote KEC005JN2025 RevA - new format full.docx')
            
            doc = Document(template_path)

            # Add additional fixture rows if needed (more than 2 fixtures)
            if len(fixtures_data) > 2:
                # Find the pricing table (it's the second table in the document)
                if len(doc.tables) > 1:
                    pricing_table = doc.tables[1]
                    add_fixture_rows_to_table(pricing_table, fixtures_data)

            # Replace text
            updated_doc = seek_and_replace(doc, replacements)

            # Save to buffer
            buffer = BytesIO()
            updated_doc.save(buffer)
            buffer.seek(0)

            # Return as download
            response = HttpResponse(
                buffer.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="Updated_Quotation_{request.POST.get("quote_no", "")}.docx"'
            return response

        except Exception as e:
            return render(request, 'quotation/index.html', {'error': str(e)})

    return render(request, 'quotation/index.html')